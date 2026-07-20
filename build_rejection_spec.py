from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('/Users/aiden/Desktop/Demo/PosDataSystem/交付文档/POS驳回功能开发规格说明.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = '165DFF'
DARK = '1D2129'
MUTED = '4E5969'
LIGHT = 'F2F4F7'
PALE_BLUE = 'E8EEF5'
PALE_RED = 'FFF1F0'
GREEN = '00A870'
RED = 'F53F3F'

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial Unicode MS'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 16, 8),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 11.5, '1F4D78', 8, 4),
]:
    st = styles[name]
    st.font.name = 'Arial Unicode MS'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ['List Bullet', 'List Number']:
    st = styles[style_name]
    st.font.name = 'Arial Unicode MS'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.28)
    st.paragraph_format.first_line_indent = Inches(-0.18)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.line_spacing = 1.2

if 'Table Note' not in styles:
    st = styles.add_style('Table Note', WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = 'Arial Unicode MS'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    st.font.size = Pt(9)
    st.font.color.rgb = RGBColor.from_string(MUTED)
    st.paragraph_format.space_before = Pt(3)
    st.paragraph_format.space_after = Pt(5)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths_dxa)
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '100')
    tblInd.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tcW = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'), str(widths_dxa[idx]))
            tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)


def add_table(headers, rows, widths, font_size=8.4, header_fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Arial Unicode MS'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
        r.font.size = Pt(font_size)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            r.font.name = 'Arial Unicode MS'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
            r.font.size = Pt(font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph('', style='Table Note').paragraph_format.space_after = Pt(1)
    return table


def add_callout(title, text, fill=LIGHT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f'{title}  ')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string(DARK)
    set_table_geometry(table, [9360])
    doc.add_paragraph('', style='Table Note').paragraph_format.space_after = Pt(1)


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p


def add_page_break():
    doc.add_page_break()


# Header / footer
header = sec.header
hp = header.paragraphs[0]
hp.text = 'POS数据管理平台｜驳回功能开发规格说明'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    run.font.name = 'Arial Unicode MS'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('内部开发交付文档  |  ')
fr.font.size = Pt(8.5)
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
fp._p.append(fld)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(42)
p.paragraph_format.space_after = Pt(8)
r = p.add_run('POS数据管理平台')
r.font.name = 'Arial Unicode MS'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
r.font.size = Pt(12)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run('驳回功能开发规格说明')
r.font.name = 'Arial Unicode MS'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
r.font.size = Pt(25)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(DARK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(26)
r = p.add_run('覆盖：单门店未匹配、质量检查-标准POS表、质量检查-异常数据')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(MUTED)

add_table(
    ['文档属性', '内容'],
    [
        ['文档类型', '产品需求 / 状态机 / 权限与接口交付规格'],
        ['版本', 'V1.0'],
        ['编制日期', '2026-07-20'],
        ['适用对象', '产品经理、前端开发、后端开发、测试、项目负责人'],
        ['实现基准', '当前项目原型及现有权限模型'],
    ],
    [2000, 7360], font_size=9.2, header_fill=LIGHT
)
add_callout('交付结论', '三个模块统一采用“POS担当驳回 → 目标营业Team处理 → POS担当复核”的责任转移闭环；状态文案统一简化，处理权限必须由角色、功能权限、数据范围、当前责任人和单据状态共同判定。')

add_page_break()
doc.add_heading('1. 文档目标与范围', level=1)
doc.add_paragraph('本文档将当前原型中的驳回功能整理为开发可执行规格，明确状态、责任人、角色操作、校验条件、数据落库、接口行为和验收标准。除特别说明外，所有规则同时适用于列表操作、单据详情操作和批量操作。')
doc.add_heading('1.1 术语', level=2)
add_table(['术语', '定义'], [
    ['POS担当', '平台管理员或具备全部数据范围的POS处理角色，负责初审、驳回、复核和最终通过。'],
    ['营业担当', '普通员工，按组织范围查看并处理其所属Team/营业所的数据。'],
    ['当前责任人', '当前必须采取下一步操作的角色或目标营业Team。'],
    ['最近操作人', '最近一次改变流程状态的实际用户。'],
    ['页面状态', '用户在列表和筛选器中看到的简化状态。'],
    ['内部状态', '后端状态机持久化使用的明确状态值。'],
], [1700, 7660], font_size=9.1)

doc.add_heading('2. 统一业务规则', level=1)
for text in [
    '驳回发起方只能是POS担当；营业担当不提供驳回按钮，也不得通过接口绕过前端发起驳回。',
    '驳回必须同时满足：POS角色、全部数据范围、模块“驳回”功能权限、单据处于允许驳回的状态。',
    '按组织范围配置的角色，其“驳回”权限在角色配置页置灰且不可选择；后端仍须重复校验。',
    '驳回后当前责任人转移至目标营业Team；营业担当只能处理其组织范围内且当前责任人确实为该Team的单据。',
    '营业担当提交处理结果后，状态转为“待复核”，当前责任人回到POS担当；营业担当不能继续编辑或重复提交。',
    'POS担当在“待复核”状态可通过或再次驳回；再次驳回重新进入“处理中”。',
    '状态、当前责任人、最近操作人必须由一次事务同时更新，禁止出现状态已变但责任人未变的中间状态。',
    '驳回针对门店级完整POS数据对象执行；异常数据按同一门店编码批量更新该门店下全部异常明细。',
]:
    bullet(text)

doc.add_heading('2.1 权限判定公式', level=2)
add_callout('驳回可执行', 'isPosActor = true AND dataScope = 全部数据 AND permission contains 模块:驳回 AND status ∈ 模块允许驳回状态。', fill=PALE_BLUE)
add_callout('营业担当可编辑', 'permission contains 模块:编辑 AND status = 营业担当处理中 AND currentOwnerType = sales AND currentOwnerTeam = 当前用户Team。', fill=PALE_BLUE)
add_callout('营业担当可提交', '满足“营业担当可编辑”全部条件，并同时具备模块“提交复核”权限。', fill=PALE_BLUE)

doc.add_heading('3. 状态机总览', level=1)
add_table(
    ['模块', '初始页面状态', '处理中', '营业提交后', 'POS最终通过后'],
    [
        ['单门店未匹配', '待处理', '处理中', '待复核', '退出未匹配列表，进入已匹配数据'],
        ['标准POS表', '待通过', '处理中', '待复核', '已通过'],
        ['异常数据', '待处理', '处理中', '待复核', '已完成，并同步标准POS表'],
    ],
    [1800, 1600, 1600, 1700, 2660], font_size=8.7
)
doc.add_paragraph('标准闭环：POS待处理/待通过 → POS驳回 → 营业处理中 → 营业提交 → POS待复核 → POS通过或再次驳回。', style='Table Note')

doc.add_heading('4. 单门店未匹配', level=1)
doc.add_paragraph('业务对象为未完成门店匹配的单门店数据。终态不保留在本列表：POS复核并完成匹配后，单据从未匹配列表移除并进入“单门店已匹配数据”。')
add_table(
    ['页面状态', '内部状态', '状态说明', '当前责任人', 'POS担当操作', '营业担当操作'],
    [
        ['待处理', '待处理', '系统初始进入未匹配池，等待POS处理', 'POS担当', '查看、编辑、校验、驳回', '组织范围内只读；不可编辑/提交/驳回'],
        ['处理中', '营业担当处理中', 'POS已驳回至指定Team', '目标营业Team', '只读；等待营业提交', '目标Team可编辑、保存、提交处理结果'],
        ['待复核', 'POS担当待处理', '营业已提交修正结果', 'POS担当', '查看、编辑、校验、再次驳回；匹配成功后转入已匹配', '只读；不可继续编辑或提交'],
    ],
    [1050, 1450, 1800, 1350, 1950, 1760], font_size=7.9
)
doc.add_heading('4.1 驳回弹窗与提交规则', level=2)
for text in [
    '展示：原始门店名称、AI异常原因、目标营业Team下拉框、手动备注。',
    '目标Team默认取单据salesTeam/team；POS担当允许在弹窗中重新选择。',
    '手动备注最大500字，当前原型允许为空；建议后端保留空值并记录AI异常原因。',
    '确认驳回后写入：rejectNote、rejectedBy、rejectedAt、currentOwnerType=sales、currentOwnerTeam、lastOperator、lastAction=驳回。',
    '营业担当提交时“处理说明”必填；成功后写入salesSubmitNote、salesSubmittedBy、salesSubmittedAt，内部状态改为POS担当待处理。',
]:
    bullet(text)
doc.add_heading('4.2 编辑与驳回组合规则', level=2)
add_table(['角色/状态', '编辑', '驳回', '说明'], [
    ['POS担当 / 待处理', '允许', '允许', 'POS可先修正，也可直接驳回。'],
    ['POS担当 / 处理中', '禁止', '禁止', '责任已转给营业，避免并发覆盖。'],
    ['POS担当 / 待复核', '允许', '允许再次驳回', '可复核修正结果。'],
    ['目标营业担当 / 处理中', '允许', '禁止', '仅处理被驳回单据。'],
    ['其他营业担当 / 任意状态', '禁止', '禁止', '即使有编辑功能权限也不可越权。'],
], [1900, 1150, 1500, 4810], font_size=8.5)

doc.add_heading('5. 质量检查-标准POS表', level=1)
doc.add_paragraph('业务对象为门店级完整标准POS表。列表可能按产品明细展示，但驳回必须以门店为单位生效，不得只驳回当前产品行。')
add_table(
    ['页面状态', '内部状态', '状态说明', '当前责任人', 'POS担当操作', '营业担当操作'],
    [
        ['待通过', '待通过', '标准POS表等待POS确认', 'POS担当', '查看、编辑、通过、驳回、删除', '组织范围内只读'],
        ['处理中', '营业担当处理中', 'POS已驳回至门店所属Team', '门店所属营业Team', '只读；等待营业提交', '目标Team可编辑并提交POS复核'],
        ['待复核', 'POS担当待处理', '营业已提交处理说明', 'POS担当', '查看、编辑、通过、再次驳回、删除', '只读'],
        ['已通过', '已通过', '质量检查完成', '-', '只读', '只读'],
    ],
    [1050, 1450, 1700, 1450, 2030, 1680], font_size=7.9
)
doc.add_heading('5.1 驳回与提交规则', level=2)
for text in [
    '目标Team固定取门店salesTeam，不允许营业担当自行认领。',
    '驳回弹窗展示门店名称、门店编码、AI判断、目标Team、手动备注；备注最大500字，当前允许为空。',
    '营业担当处理说明必填；为空时前端提示“请填写处理说明”，后端返回参数校验错误。',
    'POS通过后状态变为“已通过”，currentOwnerType=none，当前责任人显示“-”。',
    '再次驳回仅在内部状态POS担当待处理时显示，按钮文案为“再次驳回”。',
]:
    bullet(text)

doc.add_heading('6. 质量检查-异常数据', level=1)
doc.add_paragraph('业务对象为门店异常明细。一次状态操作按门店编码作用于该门店下全部异常行，确保同一门店不会出现多个责任状态。')
add_table(
    ['页面状态', '内部状态', '状态说明', '当前责任人', 'POS担当操作', '营业担当操作'],
    [
        ['待处理', '待处理', '异常待POS判断', 'POS担当', '查看、编辑、通过、驳回', '组织范围内只读'],
        ['处理中', '营业担当处理中', '已驳回至目标Team', '目标营业Team', '只读；驳回图标置灰并提示等待', '目标Team可编辑并提交POS复核'],
        ['待复核', 'POS担当待处理', '营业已提交处理结果', 'POS担当', '查看、编辑、通过、再次驳回', '只读'],
        ['已完成', '已通过', '异常处理完成并同步标准POS表', '-', '只读', '只读'],
    ],
    [1050, 1450, 1650, 1350, 2180, 1680], font_size=7.8
)
doc.add_heading('6.1 特有规则', level=2)
for text in [
    '目标Team优先依据相同storeCode在标准数据中的salesTeam；若未找到，则由ACC推导“{ACC} Team”。',
    '驳回弹窗展示AI判断、目标Team和手动备注；当前实现会回显上次驳回备注，便于再次驳回修改。',
    '营业担当提交按钮当前显示“通过”，实际语义为“提交POS复核”，提交说明当前允许为空；建议产品后续统一为必填。',
    'POS在处理中状态看见置灰驳回图标，提示“等待营业担当提交处理结果”；不能编辑、通过或再次驳回。',
    'POS最终通过后，同门店全部异常行更新为已通过，页面显示“已完成”，并同步进入标准POS表。',
]:
    bullet(text)

add_page_break()
doc.add_heading('7. 权限判定与页面控制', level=1)
add_table(
    ['控制层', '必须校验内容', '失败处理'],
    [
        ['角色配置', '按组织范围时驳回选项置灰并自动清除已选驳回权限', '提示“按组织范围的角色不能执行驳回操作”'],
        ['列表按钮', '角色、功能权限、数据范围、内部状态、当前责任人', '隐藏按钮；处理中可保留置灰提示'],
        ['弹窗提交', '重新获取最新状态并复核权限', '关闭提交或提示“当前账号或单据状态不可驳回”'],
        ['后端接口', 'Token角色、数据范围、Team、状态版本、功能权限', '403权限不足；409状态冲突；422参数错误'],
        ['数据查询', '营业担当仅返回组织范围内数据', '后端过滤，不能只依赖前端隐藏'],
    ],
    [1500, 5360, 2500], font_size=8.6
)
doc.add_heading('7.1 页面按钮矩阵', level=2)
add_table(['角色', '待处理/待通过', '处理中', '待复核', '终态'], [
    ['POS担当', '编辑；通过（质量检查）；驳回', '只读等待', '编辑；通过；再次驳回', '只读'],
    ['目标营业担当', '只读', '编辑；保存；提交POS复核', '只读', '只读'],
    ['非目标营业担当', '只读或不可见', '只读或不可见', '只读或不可见', '只读或不可见'],
], [1650, 2200, 2100, 2100, 1310], font_size=8.5)

doc.add_heading('8. 数据模型与接口建议', level=1)
doc.add_heading('8.1 工作流字段', level=2)
add_table(['字段', '类型', '必填时机', '说明'], [
    ['workflowStatus', 'enum', '始终', '内部状态：待处理/待通过、营业担当处理中、POS担当待处理、已通过。'],
    ['currentOwnerType', 'enum', '始终', 'pos / sales / none。'],
    ['currentOwnerName', 'string', '始终', 'POS担当、目标Team或“-”。'],
    ['currentOwnerTeam', 'string', '营业处理中', '目标营业Team。'],
    ['rejectNote', 'string(500)', '驳回时可选', 'POS手动备注。'],
    ['rejectedBy / rejectedAt', 'string/datetime', '驳回时', '驳回操作人及时间。'],
    ['salesSubmitNote', 'string', '营业提交时', '未匹配和标准POS表必填；异常数据当前可选。'],
    ['salesSubmittedBy / At', 'string/datetime', '营业提交时', '营业处理提交人及时间。'],
    ['lastOperatorName / Role', 'string', '每次状态变化', '最近操作人及角色。'],
    ['lastAction / lastOperatedAt', 'string/datetime', '每次状态变化', '驳回、提交复核、通过等。'],
    ['version', 'integer', '始终', '乐观锁版本，防止重复提交和并发覆盖。'],
], [2200, 1500, 1850, 3810], font_size=8.2)

doc.add_heading('8.2 建议接口', level=2)
add_table(['接口', '方法', '核心请求参数', '成功后的状态'], [
    ['/workflow/{module}/{documentId}/reject', 'POST', 'targetTeam, rejectNote, version', '营业担当处理中'],
    ['/workflow/{module}/{documentId}/sales-submit', 'POST', 'salesSubmitNote, version', 'POS担当待处理'],
    ['/workflow/{module}/{documentId}/approve', 'POST', 'version', '标准POS=已通过；异常=已通过；未匹配=转入已匹配'],
    ['/workflow/{module}/{documentId}', 'GET', '-', '返回状态、责任人、操作记录和可执行动作'],
], [3000, 900, 3200, 2260], font_size=8.2)
add_callout('接口原则', '后端返回allowedActions数组，由服务端计算最终可操作项；前端按钮权限只用于交互体验，不作为安全边界。')

doc.add_heading('9. 异常处理与并发控制', level=1)
for text in [
    '重复点击：同一version的重复驳回/提交/通过请求只允许一次成功，其余返回409。',
    '状态过期：弹窗打开后若其他用户已改变状态，提交时返回409并刷新列表。',
    '组织变更：用户Team变更后，不得继续处理原Team单据；后端按请求时组织关系校验。',
    '目标Team为空：禁止驳回并提示“未配置目标营业Team”。',
    '异常数据门店级事务：同一storeCode下全部明细更新必须原子提交。',
    '未匹配转入已匹配：匹配成功、源列表移除、目标列表新增和日志写入必须在同一事务或可靠事件中完成。',
    '所有状态变更写审计日志，至少记录beforeStatus、afterStatus、operator、role、targetTeam、note、timestamp、requestId。',
]:
    bullet(text)

doc.add_heading('10. 验收用例', level=1)
add_table(['编号', '场景', '前置条件', '操作', '预期结果'], [
    ['AC-01', 'POS首次驳回', '允许驳回的初始状态', '选择Team并确认驳回', '状态=处理中；责任人=目标Team；最近操作人=POS用户。'],
    ['AC-02', '营业编辑限制', '状态=待处理', '营业点击编辑', '按钮不可用，提示需先由POS驳回。'],
    ['AC-03', '目标Team处理', '状态=处理中且Team匹配', '编辑并提交处理结果', '状态=待复核；责任人=POS担当；最近操作人=营业用户。'],
    ['AC-04', '非目标Team越权', '状态=处理中但Team不匹配', '调用编辑或提交接口', '前端无按钮；后端返回403。'],
    ['AC-05', 'POS处理中限制', '状态=处理中', '尝试编辑/通过/再次驳回', '操作不可执行；异常数据驳回图标置灰。'],
    ['AC-06', 'POS再次驳回', '状态=待复核', '点击再次驳回', '状态回到处理中；责任人转为目标Team；保留完整历史。'],
    ['AC-07', '标准POS最终通过', '标准POS状态=待复核', 'POS点击通过', '状态=已通过；责任人=-；最近操作人=POS用户。'],
    ['AC-08', '异常最终通过', '异常状态=待复核', 'POS点击通过', '同门店异常均为已通过，页面显示已完成并同步标准POS。'],
    ['AC-09', '未匹配复核成功', '未匹配状态=待复核', 'POS校验匹配成功', '单据退出未匹配列表并进入已匹配列表。'],
    ['AC-10', '按组织范围配置驳回', '角色数据范围=按组织范围', '打开角色权限', '驳回复选框置灰且不可保存为已选。'],
    ['AC-11', '并发状态冲突', '两个用户持有同一version', '先后提交状态变更', '首个成功，第二个409并刷新。'],
    ['AC-12', '详情字段', '任意流程状态', '打开单据详情', '显示当前责任人、最近操作人、驳回原因、处理说明及状态日志。'],
], [750, 1400, 1900, 1850, 3460], font_size=7.7)

doc.add_heading('10.1 开发完成定义（DoD）', level=2)
for text in [
    '三个模块均完成前端状态、按钮、弹窗、筛选器和详情字段改造。',
    '后端完成统一权限校验、状态机、事务、乐观锁与审计日志。',
    'POS账号与营业担当账号可完成至少一次“驳回—处理—复核—通过”双账号闭环。',
    '上述12条验收用例全部通过，并覆盖接口级越权测试。',
    '产品确认三个模块的终态文案差异：未匹配转列表、标准POS“已通过”、异常数据“已完成”。',
]:
    bullet(text)

doc.add_heading('附录A 状态转换校验表', level=2)
add_table(['当前内部状态', '允许动作', '目标内部状态', '责任人变化'], [
    ['待处理 / 待通过', 'POS驳回', '营业担当处理中', 'POS担当 → 目标营业Team'],
    ['待处理 / 待通过', 'POS通过（质量检查）', '已通过', 'POS担当 → -'],
    ['营业担当处理中', '营业提交复核', 'POS担当待处理', '目标营业Team → POS担当'],
    ['POS担当待处理', 'POS再次驳回', '营业担当处理中', 'POS担当 → 目标营业Team'],
    ['POS担当待处理', 'POS通过', '已通过或转入已匹配', 'POS担当 → - / 流程结束'],
], [2200, 1900, 2400, 2860], font_size=8.4)
doc.add_paragraph('任何未在上表中的状态转换均应由后端拒绝，返回409状态冲突或422非法状态转换。', style='Table Note')

doc.core_properties.title = 'POS驳回功能开发规格说明'
doc.core_properties.subject = '单门店未匹配、标准POS表、异常数据驳回闭环'
doc.core_properties.author = 'POS数据管理平台项目组'
doc.core_properties.keywords = '驳回, 状态机, 权限, POS担当, 营业担当'
doc.save(OUT)
print(OUT)
