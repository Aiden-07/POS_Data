from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "PRD-台账与汇总及系统设置.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_doc_font(doc):
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 41, 55)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(75, 85, 99)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, True)
        set_cell_text(table.cell(i, 1), v)
        set_cell_shading(table.cell(i, 0), "EEF4FF")
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], "EAF1FF")
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    set_doc_font(doc)

    add_title(doc, "POS数据管理平台 PRD", "台账与汇总、系统设置")
    add_kv_table(
        doc,
        [
            ("文档版本", "V1.0"),
            ("适用范围", "台账与汇总、系统设置，包括用户管理、角色管理、系统日志及字段配置能力说明"),
            ("目标读者", "产品经理、前端工程师、后端工程师、测试工程师、项目负责人"),
            ("交付目标", "明确页面信息架构、核心交互、接口、数据模型与验收标准，可直接用于研发拆分和测试验收。"),
        ],
    )

    doc.add_heading("1. 文档概述", level=1)
    doc.add_paragraph(
        "本文档定义 POS 数据管理平台中“台账与汇总”和“系统设置”两个目录的产品需求。"
        "其中台账与汇总承接质量检查后的标准 POS 明细数据，支持筛选、表头配置、分组、导出和单据详情追踪；"
        "系统设置负责用户、角色、权限、日志与字段标准配置，保障平台的功能权限、数据权限和审计追溯能力。"
    )

    doc.add_heading("2. 信息架构与导航", level=1)
    add_table(
        doc,
        ["一级目录", "二级目录/页面", "定位", "主要能力"],
        [
            ["台账与汇总", "标准 POS 明细台账", "标准化明细数据查询与导出", "筛选、表头、分组、导出、详情、数据条数展示"],
            ["系统设置", "用户管理", "账号与角色绑定", "新建用户、编辑、停用、删除、重置密码、继承角色权限"],
            ["系统设置", "角色管理", "功能权限与数据权限配置", "角色新增/编辑、组织范围选择、功能权限勾选、删除"],
            ["系统设置", "系统日志", "平台操作审计", "登录登出、查看、编辑、通过、驳回、导出、配置变更记录"],
            ["系统设置", "字段配置能力", "字段标准管理能力", "用于质量检查和台账表头，当前不作为系统设置二级目录展示"],
        ],
    )

    doc.add_heading("3. 台账与汇总", level=1)
    doc.add_heading("3.1 页面目标", level=2)
    add_bullets(
        doc,
        [
            "展示通过质量检查并进入标准化流程的 POS 明细数据。",
            "支持按时间、组织、经销商、门店、产品维度检索和导出。",
            "支持表头字段配置和分组查看，提升运营人员分析效率。",
            "支持从单据入口查看来源邮件、来源附件和状态流转日志。",
        ],
    )

    doc.add_heading("3.2 筛选区", level=2)
    add_table(
        doc,
        ["控件", "默认值", "交互规则", "研发说明"],
        [
            ["年份", "2026年", "下拉选择年份，影响台账明细查询范围", "字段：year"],
            ["月份", "6月", "下拉选择月份，默认展示当前业务月", "字段：month"],
            ["关键字", "空", "支持经销商、门店名称/编码、产品名称/编码模糊搜索", "字段：keyword"],
            ["所属组织", "全部", "按区域、营业所、经销商逐级筛选", "字段：region、office、dealer"],
            ["查询", "-", "按当前条件刷新列表", "调用 GET /api/ledger/records"],
            ["重置", "-", "恢复默认条件并刷新列表", "重置本地筛选状态"],
        ],
    )

    doc.add_heading("3.3 台账列表字段", level=2)
    add_table(
        doc,
        ["字段", "说明", "是否默认展示", "备注"],
        [
            ["年月", "数据所属年月", "是", "格式：YYYY年MM月"],
            ["ACC", "账号或渠道分类字段", "是", "来源于标准 POS 字段"],
            ["经销商名称", "经销商主体名称", "是", "组织权限过滤依据之一"],
            ["门店编码", "门店唯一编码", "是", "支持搜索"],
            ["门店名称", "门店名称", "是", "支持搜索"],
            ["产品编码", "产品编码", "是", "原 69 码列已调整为产品编码相关字段"],
            ["产品名称", "产品名称", "是", "支持搜索"],
            ["69码", "商品条码", "是", "保留标准字段"],
            ["销售数量", "销售数量", "是", "数值字段"],
            ["销售金额", "销售金额", "是", "金额字段"],
            ["销售成本", "销售成本", "是", "金额字段"],
            ["零售价", "零售价", "是", "金额字段"],
            ["区域", "组织区域", "可选", "选中后固定展示在经销商名称前"],
            ["营业所", "二级组织", "可选", "选中后固定展示在经销商名称前"],
        ],
    )

    doc.add_heading("3.4 表头配置", level=2)
    add_bullets(
        doc,
        [
            "列表右上角提供“表头”按钮，点击后展示可勾选字段面板。",
            "标准字段默认选中，区域和营业所为可选字段。",
            "区域、营业所被勾选后，展示位置固定在经销商名称前，不随勾选顺序变化。",
            "取消勾选后，对应列立即从列表隐藏，不影响数据本身。",
        ],
    )

    doc.add_heading("3.5 分组与导出", level=2)
    add_table(
        doc,
        ["能力", "交互", "规则"],
        [
            ["分组", "点击分组按钮后选择区域、营业所或经销商", "单选；选择后列表按字段聚合，组支持折叠/展开"],
            ["导出", "点击导出按钮", "导出当前筛选、当前表头范围内的数据"],
            ["数据条数", "列表右下角展示", "展示当前筛选后的单据数量，例如“当前 480 条单据”"],
        ],
    )

    doc.add_heading("3.6 单据详情", level=2)
    add_bullets(
        doc,
        [
            "标准 POS 门店列表的操作列新增“详情”入口，不替换门店名称点击预览逻辑。",
            "详情页包含来源邮件、来源附件、当前模块、当前节点、状态流转日志。",
            "来源附件需要保留附件结构；如果来源为压缩包，支持展开查看压缩包内文件并预览对应内容。",
            "详情页用于追踪数据从文件收取、质量检查到台账与汇总的完整链路。",
        ],
    )

    doc.add_heading("4. 系统设置", level=1)
    doc.add_heading("4.1 用户管理", level=2)
    add_table(
        doc,
        ["功能", "需求描述", "交互规则"],
        [
            ["用户列表", "展示用户姓名、登录账号、所属角色、继承数据权限、状态、最近登录、操作", "操作列左对齐"],
            ["新建用户", "支持设置姓名、账号、密码、所属角色", "用户不单独配置数据权限"],
            ["权限继承", "用户继承角色下的功能权限和数据权限", "修改角色后关联用户自动生效"],
            ["操作", "编辑、停用、重置密码、删除", "删除需二次确认"],
            ["搜索筛选", "按用户姓名、账号、手机号搜索，按状态和角色筛选", "实时或点击查询刷新"],
        ],
    )

    doc.add_heading("4.2 角色管理", level=2)
    add_table(
        doc,
        ["区域", "需求描述", "交互规则"],
        [
            ["角色列表", "展示角色名称、角色说明、功能权限、角色数据权限、关联用户、操作", "状态列、复制、停用已移除"],
            ["新建/编辑角色", "配置角色名称、数据权限、功能权限", "角色说明已移除"],
            ["数据权限", "合并原“数据权限策略”和“数据权限范围”", "模式包括全部数据、按组织范围"],
            ["组织范围", "下拉树组件展示区域和营业所", "选择一级区域后才展示二级营业所"],
            ["复选框", "每个区域和营业所前均有复选框", "支持全选、半选、取消选择"],
            ["帮助说明", "数据权限标题右侧提供问号", "hover 展示说明弹层，不占用弹窗内容区域"],
        ],
    )

    doc.add_heading("4.3 组织范围下拉树", level=2)
    add_table(
        doc,
        ["状态", "展示", "规则"],
        [
            ["未选择", "输入框占位：请选择组织范围", "下拉面板默认展示一级区域"],
            ["选择区域", "输入框展示区域标签", "右侧二级面板展示该区域下营业所"],
            ["选择营业所", "输入框展示区域/营业所标签，可折叠标签", "数据权限精确到选中营业所"],
            ["全选", "选择全部区域和营业所", "等价于全部数据或全部组织范围"],
            ["取消", "取消勾选后同步更新标签和权限范围", "父子级联动"],
        ],
    )

    doc.add_heading("4.4 功能权限", level=2)
    add_table(
        doc,
        ["模块", "可配置动作"],
        [
            ["首页概览", "查看"],
            ["文件收取", "查看、新建、编辑、导出"],
            ["质量检查", "查看、编辑、审批、导出"],
            ["台账与汇总", "查看、导出"],
            ["数据分析", "查看、导出"],
            ["系统设置", "查看、新建、编辑、删除"],
        ],
    )

    doc.add_heading("4.5 系统日志", level=2)
    add_bullets(
        doc,
        [
            "记录用户登录、登出、会话超时等上下线行为。",
            "记录用户在平台内的查看、编辑、修改、通过、驳回、覆盖、忽略、导出等操作。",
            "质量检查中修改某条异常数据时，应记录对象、原值、新值、操作人、操作时间和结果。",
            "支持按时间、用户、模块、操作类型、对象 ID、操作结果筛选。",
            "日志数据不可由普通用户编辑或删除，仅允许有权限用户查询和导出。",
        ],
    )

    doc.add_heading("4.6 字段配置能力", level=2)
    doc.add_paragraph(
        "字段配置用于定义质量检查和台账表头的标准字段。当前二级目录中不展示“字段配置”，但能力可作为后台配置或后续入口开放。"
        "字段配置影响质量检查对原始 POS 表字段的校验，也影响台账与汇总中的字段表头展示。区域、营业所由组织架构生成，不受字段配置控制。"
    )
    add_table(
        doc,
        ["标准字段", "字段编码", "影响范围"],
        [
            ["年月", "month", "质量检查、台账表头"],
            ["ACC", "acc", "质量检查、台账表头"],
            ["经销商名称", "dealerName", "质量检查、台账表头"],
            ["门店编码", "storeCode", "质量检查、台账表头"],
            ["门店名称", "storeName", "质量检查、台账表头"],
            ["产品编码", "productCode", "质量检查、台账表头"],
            ["产品名称", "productName", "质量检查、台账表头"],
            ["69码", "barcode69", "质量检查、台账表头"],
            ["销售数量", "quantity", "质量检查、台账表头"],
            ["销售金额", "amount", "质量检查、台账表头"],
            ["销售成本", "cost", "质量检查、台账表头"],
            ["零售价", "retailPrice", "质量检查、台账表头"],
        ],
    )

    doc.add_heading("5. 接口设计", level=1)
    add_table(
        doc,
        ["接口", "方法", "用途", "关键入参", "返回"],
        [
            ["/api/ledger/records", "GET", "查询台账明细", "year、month、keyword、region、office、dealer、pageSize", "records、total"],
            ["/api/ledger/records/{id}", "GET", "查询单据详情", "id", "record、sourceMail、attachments、statusLogs"],
            ["/api/ledger/export", "POST", "导出当前台账数据", "filters、columns、groupBy", "downloadUrl"],
            ["/api/settings/users", "GET", "用户列表", "keyword、roleId、status", "users、total"],
            ["/api/settings/users", "POST", "新建用户", "name、account、password、roleId", "userId"],
            ["/api/settings/users/{id}", "PUT", "编辑用户", "name、roleId、status", "success"],
            ["/api/settings/users/{id}", "DELETE", "删除用户", "id", "success"],
            ["/api/settings/roles", "GET", "角色列表", "keyword", "roles、total"],
            ["/api/settings/roles", "POST", "新建角色", "name、dataScope、permissions", "roleId"],
            ["/api/settings/roles/{id}", "PUT", "编辑角色", "name、dataScope、permissions", "success"],
            ["/api/settings/roles/{id}", "DELETE", "删除角色", "id", "success"],
            ["/api/settings/logs", "GET", "查询系统日志", "timeRange、userId、module、action、result", "logs、total"],
        ],
    )

    doc.add_heading("6. 数据模型", level=1)
    add_table(
        doc,
        ["模型", "关键字段", "说明"],
        [
            ["ledger_record", "id、month、acc、dealerName、storeCode、storeName、productCode、productName、barcode69、quantity、amount、cost、retailPrice、region、office", "标准 POS 明细台账"],
            ["user", "id、name、account、phone、roleId、status、lastLoginAt", "平台用户"],
            ["role", "id、name、dataScopeType、status、createdAt、updatedAt", "角色主表"],
            ["role_permission", "roleId、module、actions", "角色功能权限"],
            ["org_scope", "roleId、regionCode、officeCode", "角色组织数据权限"],
            ["system_log", "id、userId、module、action、objectType、objectId、beforeValue、afterValue、result、createdAt", "系统操作日志"],
            ["field_config", "id、fieldName、fieldCode、qaEnabled、ledgerVisible、status、sortNo", "标准字段配置"],
        ],
    )

    doc.add_heading("7. 验收用例", level=1)
    add_table(
        doc,
        ["编号", "场景", "前置条件", "操作", "预期结果"],
        [
            ["AC-001", "台账查询", "存在 2026年06月数据", "选择年份月份并点击查询", "列表仅展示对应年月数据，右下角数据条数正确"],
            ["AC-002", "关键字搜索", "存在门店、产品、经销商数据", "输入门店编码或产品名称搜索", "返回匹配记录"],
            ["AC-003", "表头配置", "进入台账页", "勾选区域、营业所", "两列固定展示在经销商名称前"],
            ["AC-004", "台账分组", "进入台账页", "按区域分组", "展示区域分组行，支持折叠展开"],
            ["AC-005", "台账导出", "存在筛选条件", "点击导出", "导出当前条件下数据"],
            ["AC-006", "用户继承角色权限", "用户绑定华北区域经理", "登录用户查看数据", "仅可见华北区域权限范围内数据"],
            ["AC-007", "角色组织范围选择", "打开新建角色弹窗", "选择华北区域后选择石家庄营业所", "输入框展示已选标签，权限范围保存正确"],
            ["AC-008", "角色删除", "角色无不可删除限制", "点击删除并确认", "角色从列表移除并记录系统日志"],
            ["AC-009", "系统日志", "用户执行驳回操作", "进入系统日志查询", "可查到操作人、对象、时间、结果和详情"],
            ["AC-010", "字段配置影响", "修改标准字段显示状态", "进入台账与质量检查", "对应字段校验和表头展示按配置生效"],
        ],
    )

    doc.add_heading("8. 非功能要求", level=1)
    add_bullets(
        doc,
        [
            "列表查询响应时间：常规筛选应在 2 秒内返回，导出任务可异步执行。",
            "权限控制：前端隐藏无权限入口，后端必须二次校验功能权限和数据权限。",
            "审计追踪：涉及数据变更、权限变更和导出的操作必须写入系统日志。",
            "可用性：表格横向滚动、下拉树和弹窗在常见桌面分辨率下不得出现内容遮挡。",
            "安全性：密码不明文存储；重置密码、删除用户、删除角色需要二次确认。",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
