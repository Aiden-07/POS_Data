from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/aiden/Desktop/Demo/PosDataSystem/PRD-台账与汇总及系统设置.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "EAF2FF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(17, 24, 39)
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
    return p


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("POS数据管理平台 PRD：台账与汇总、系统设置")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)

    add_para(doc, "版本：V1.0    日期：2026-07-06    适用对象：产品、前端、后端、测试、实施")

    add_heading(doc, "一、文档目标", 1)
    add_para(
        doc,
        "本文档用于明确 POS 数据管理平台中“台账与汇总”和“系统设置”两个目录的功能范围、页面交互、数据规则、接口需求、数据模型与验收标准，作为研发实施和测试验收依据。",
    )

    add_heading(doc, "二、范围说明", 1)
    add_table(
        doc,
        ["一级目录", "模块/页面", "本期范围", "说明"],
        [
            ["台账与汇总", "标准 POS 明细台账", "包含", "承载标准化 POS 明细数据的查询、分组、导出、详情查看。"],
            ["台账与汇总", "筛选模块", "包含", "支持时间、关键词、组织范围等条件组合查询。"],
            ["台账与汇总", "表头配置", "包含", "控制列表字段显示，区域、营业所可选并固定插入到经销商名称前。"],
            ["台账与汇总", "分组展示", "包含", "支持按区域、营业所、经销商单选分组，并支持展开/折叠。"],
            ["系统设置", "用户管理", "包含", "用户账号、角色分配、状态、重置密码、删除等管理能力。"],
            ["系统设置", "角色管理", "包含", "角色功能权限和数据权限统一配置。"],
            ["系统设置", "系统日志", "包含", "记录用户上下线及平台内关键操作审计。"],
            ["系统设置", "字段配置", "不作为菜单展示", "字段能力由配置支撑，字段配置二级目录在导航中隐藏。"],
        ],
    )

    add_heading(doc, "三、信息架构与布局交互", 1)
    add_table(
        doc,
        ["页面", "区域", "布局", "核心交互", "研发说明"],
        [
            ["台账与汇总", "筛选区", "顶部独立筛选模块", "默认展示常用筛选；右侧支持展开更多筛选条件。", "查询条件变更后需刷新当前台账结果与记录数。"],
            ["台账与汇总", "工具栏", "列表上方左右布局", "左侧表头配置，右侧分组与导出。", "导出应基于当前筛选与当前可见字段。"],
            ["台账与汇总", "明细列表", "横向可滚动表格", "支持字段显示控制、分组折叠、详情入口。", "列表底部右下角显示当前单据数量。"],
            ["系统设置-用户管理", "筛选区", "表格上方", "按关键词、状态、角色筛选用户。", "用户不再单独配置数据权限。"],
            ["系统设置-用户管理", "列表区", "表格", "编辑、停用、重置密码、删除。", "用户权限继承所属角色的功能权限和数据权限。"],
            ["系统设置-角色管理", "角色列表", "表格", "新建、编辑、删除。", "移除复制、停用、状态字段。"],
            ["系统设置-角色管理", "角色弹窗", "单列配置区 + 权限矩阵", "配置角色名称、数据权限、功能权限。", "数据权限策略与范围合并为一个组织范围配置。"],
            ["系统设置-日志", "日志列表", "筛选 + 表格", "按用户、模块、动作、时间、结果筛选。", "日志不可编辑，仅支持查询和导出。"],
        ],
    )

    add_heading(doc, "四、台账与汇总功能需求", 1)
    add_heading(doc, "4.1 标准 POS 明细台账", 2)
    add_para(doc, "用于展示已通过质量检查并进入标准 POS 明细的数据，是后续统计分析和导出的核心数据源。")
    add_table(
        doc,
        ["字段", "类型", "展示规则", "交互规则"],
        [
            ["年月", "文本", "格式为 YYYY年MM月", "随筛选时间联动。"],
            ["ACC", "文本", "展示客户/渠道归类", "支持表头配置显示/隐藏。"],
            ["经销商名称", "文本", "展示经销商名称", "当区域、营业所被选中显示时，二者固定展示在该字段前。"],
            ["门店编码", "文本", "展示标准门店编码", "支持关键词搜索。"],
            ["门店名称", "文本", "展示标准门店名称", "支持关键词搜索。"],
            ["产品编码", "文本", "展示产品编码", "支持关键词搜索。"],
            ["产品名称", "文本", "展示产品名称", "支持关键词搜索。"],
            ["69码", "文本", "展示商品 69 码", "支持表头配置显示/隐藏。"],
            ["销售数量", "数值", "展示销售数量", "默认显示。"],
            ["销售金额", "金额", "展示销售金额", "默认显示。"],
            ["销售成本", "金额", "展示销售成本", "默认显示。"],
            ["零售价", "金额", "展示零售价", "默认显示。"],
        ],
    )

    add_heading(doc, "4.2 筛选模块", 2)
    add_table(
        doc,
        ["筛选项", "默认值", "交互说明", "查询影响"],
        [
            ["选择时间", "2026年6月", "年、月下拉选择。", "过滤对应年月明细。"],
            ["关键字", "空", "单输入框，支持经销商、门店名称/编码、产品名称/编码搜索。", "模糊匹配明细记录。"],
            ["所属组织", "全部", "默认展示区域层级；选择一级区域后展示其下属营业所。", "过滤组织权限范围内数据。"],
            ["更多筛选", "收起", "点击右侧展开其它筛选条件。", "展开后不影响默认查询，只有用户选择后生效。"],
            ["查询", "-", "执行筛选。", "刷新列表与记录数。"],
            ["重置", "-", "清空筛选并恢复默认值。", "刷新为默认数据。"],
        ],
    )

    add_heading(doc, "4.3 表头配置", 2)
    add_bullets(
        doc,
        [
            "入口位于台账列表左上方，名称为“表头”。",
            "默认选中当前列表基础字段。",
            "可选字段包含“区域”“营业所”，不包含经销商扩展项。",
            "区域、营业所被选中后，固定插入到“经销商名称”字段前方。",
            "字段勾选/取消后，列表表头和对应列实时变化。",
        ],
    )

    add_heading(doc, "4.4 分组与导出", 2)
    add_table(
        doc,
        ["功能", "入口", "规则", "验收点"],
        [
            ["分组", "列表右上角", "支持按区域、营业所、经销商单选分组。", "选择后列表按分组行展示，并支持展开/折叠。"],
            ["导出", "分组按钮右侧", "导出当前筛选结果和当前可见字段。", "导出文件字段顺序与页面可见表头一致。"],
            ["记录数", "列表右下角", "展示当前筛选后单据数量。", "筛选变化后数量同步变化。"],
        ],
    )

    add_heading(doc, "4.5 单据详情", 2)
    add_bullets(
        doc,
        [
            "台账明细操作列新增“详情”图标入口。",
            "详情页不替代门店名称点击预览逻辑。",
            "详情页展示来源邮件、来源附件、附件结构、状态流转日志。",
            "若来源附件为压缩包，保留压缩包结构，支持展开内部文件并预览文件内容。",
            "状态日志需记录从文件收取、质量检查到台账汇总的完整流转。",
        ],
    )

    add_heading(doc, "五、系统设置功能需求", 1)
    add_heading(doc, "5.1 用户管理", 2)
    add_table(
        doc,
        ["功能", "规则", "交互", "备注"],
        [
            ["用户列表", "展示用户姓名、登录账号、所属角色、继承数据权限、状态、最近登录、操作。", "支持关键词、状态、角色筛选。", "操作列左对齐。"],
            ["新建用户", "填写用户姓名、账号、手机号、初始密码、角色。", "保存后用户继承角色权限。", "不配置用户级数据权限。"],
            ["编辑用户", "可修改基础信息和所属角色。", "角色变化后继承权限同步变化。", "需记录系统日志。"],
            ["停用用户", "停用后不可登录。", "二次确认。", "需记录系统日志。"],
            ["重置密码", "管理员可重置用户密码。", "弹窗确认。", "需记录系统日志。"],
            ["删除用户", "删除用户账号。", "二次确认。", "默认示例数据可模拟删除。"],
        ],
    )

    add_heading(doc, "5.2 角色管理", 2)
    add_para(doc, "角色管理负责统一维护功能权限和数据权限。用户只分配角色，不再单独维护数据权限。")
    add_table(
        doc,
        ["配置项", "规则", "交互说明"],
        [
            ["角色名称", "必填，唯一。", "单行输入。"],
            ["数据权限", "合并原“数据权限策略”和“数据权限范围”。", "选择全部数据或按组织范围。"],
            ["组织范围", "按区域和营业所两级结构选择。", "选择一级区域后才显示其二级营业所；支持复选和全选。"],
            ["功能权限", "按模块配置查看、新建、编辑、审核、导出、删除等动作。", "勾选即生效，保存角色后更新权限。"],
            ["删除角色", "删除前需校验是否有关联用户。", "有关联用户时提示先调整用户角色。"],
        ],
    )
    add_table(
        doc,
        ["权限模式", "说明", "数据可见范围"],
        [
            ["全部数据", "适用于系统管理员。", "可查看全部区域、营业所、经销商、门店和 POS 明细。"],
            ["按组织范围", "适用于区域经理、营业所负责人。", "根据勾选的区域/营业所限制数据范围。"],
        ],
    )

    add_heading(doc, "5.3 系统日志", 2)
    add_table(
        doc,
        ["日志类型", "触发动作", "记录字段", "用途"],
        [
            ["登录日志", "登录成功、登录失败、退出登录。", "用户、账号、时间、IP、结果、失败原因。", "安全审计。"],
            ["数据操作日志", "查看、编辑、修改、驳回、通过、导出。", "模块、单据ID、操作前后状态、操作人、时间。", "追踪数据流转。"],
            ["权限操作日志", "新建/编辑/删除用户、角色，重置密码。", "对象、变更内容、操作人、时间。", "权限审计。"],
            ["质量检查日志", "异常修正、驳回、审核通过。", "异常数据ID、字段、原值、新值、原因。", "质检追溯。"],
        ],
    )

    add_heading(doc, "5.4 字段配置说明", 2)
    add_bullets(
        doc,
        [
            "系统设置左侧二级目录不展示“字段配置”。",
            "标准字段仍作为后台配置能力存在，影响质量检查字段校验与台账表头字段展示。",
            "标准字段包括：年月、ACC、经销商名称、门店编码、门店名称、产品编码、产品名称、69码、销售数量、销售金额、销售成本、零售价。",
            "区域、营业所为组织结构生成字段，不受字段配置控制。",
        ],
    )

    add_heading(doc, "六、状态与权限规则", 1)
    add_table(
        doc,
        ["规则", "说明", "验收标准"],
        [
            ["用户权限继承", "用户只继承所属角色的功能权限和数据权限。", "用户管理中不出现独立数据权限配置。"],
            ["数据权限优先级", "角色选择全部数据时不限制组织；选择按组织范围时必须至少选择一个区域或营业所。", "未选择组织范围时不可保存。"],
            ["组织层级", "区域为一级，营业所为二级。", "未选择区域时不展示该区域营业所。"],
            ["审计留痕", "所有新增、编辑、删除、导出、审核动作写入系统日志。", "日志列表可查询到记录。"],
        ],
    )

    add_heading(doc, "七、研发接口", 1)
    add_table(
        doc,
        ["接口", "方法", "说明", "关键参数/返回"],
        [
            ["/api/ledger/records", "GET", "查询台账明细", "year, month, keyword, orgScope, groupBy, visibleFields"],
            ["/api/ledger/records/{id}", "GET", "查询单据详情", "recordId, sourceMail, sourceAttachments, flowLogs"],
            ["/api/ledger/records/{id}/flow-logs", "GET", "查询流转日志", "recordId, node, action, operator, time"],
            ["/api/ledger/export", "POST", "导出台账数据", "filters, visibleFields"],
            ["/api/settings/users", "GET", "查询用户列表", "keyword, status, roleId"],
            ["/api/settings/users", "POST", "新建用户", "name, account, phone, password, roleId"],
            ["/api/settings/users/{id}", "PUT", "编辑用户", "name, phone, roleId, status"],
            ["/api/settings/users/{id}", "DELETE", "删除用户", "userId"],
            ["/api/settings/users/{id}/reset-password", "POST", "重置密码", "userId, newPassword"],
            ["/api/settings/roles", "GET", "查询角色列表", "keyword"],
            ["/api/settings/roles", "POST", "新建角色", "roleName, dataScope, functionPermissions"],
            ["/api/settings/roles/{id}", "PUT", "编辑角色", "roleId, roleName, dataScope, functionPermissions"],
            ["/api/settings/roles/{id}", "DELETE", "删除角色", "roleId"],
            ["/api/settings/org-tree", "GET", "查询组织树", "regions, offices"],
            ["/api/settings/logs", "GET", "查询系统日志", "operator, module, action, dateRange, result"],
        ],
    )

    add_heading(doc, "八、数据模型", 1)
    add_table(
        doc,
        ["表名", "用途", "核心字段"],
        [
            ["pos_ledger_record", "台账明细主表", "id, month, acc, dealer_name, store_code, store_name, product_code, product_name, barcode69, qty, amount, cost, retail_price, region, office"],
            ["document_source_mail", "来源邮件", "id, subject, content, provider, provided_at"],
            ["document_source_attachment", "来源附件", "id, mail_id, file_name, file_type, status, parent_zip_id"],
            ["document_flow_log", "单据流转日志", "id, record_id, node, action, from_status, to_status, operator, operated_at, remark"],
            ["sys_user", "用户表", "id, name, account, phone, role_id, status, last_login_at"],
            ["sys_role", "角色表", "id, role_name, data_scope_type, created_at, updated_at"],
            ["sys_role_data_scope", "角色数据权限", "id, role_id, region_code, office_code"],
            ["sys_role_function_permission", "角色功能权限", "id, role_id, module, action"],
            ["sys_operation_log", "系统操作日志", "id, operator, module, action, target_id, before_value, after_value, result, operated_at"],
        ],
    )

    add_heading(doc, "九、验收用例", 1)
    add_table(
        doc,
        ["编号", "场景", "操作步骤", "预期结果"],
        [
            ["TC-001", "台账筛选", "选择 2026年6月并输入门店关键字后点击查询。", "列表仅展示匹配数据，右下角数量同步更新。"],
            ["TC-002", "表头配置", "勾选区域、营业所。", "列表新增两列，且固定显示在经销商名称前。"],
            ["TC-003", "分组展示", "选择按区域分组。", "列表出现区域分组行，可展开/折叠明细。"],
            ["TC-004", "导出", "筛选后点击导出。", "导出文件与当前筛选结果和可见字段一致。"],
            ["TC-005", "用户继承权限", "创建用户并分配华北区域经理角色。", "用户继承该角色功能权限和华北区域数据权限。"],
            ["TC-006", "角色组织范围", "选择按组织范围，勾选华北区域后选择石家庄营业所。", "保存后角色仅可访问对应组织数据。"],
            ["TC-007", "删除角色", "删除有关联用户的角色。", "系统提示存在关联用户，不允许删除。"],
            ["TC-008", "系统日志", "执行用户重置密码、台账导出。", "系统日志生成对应记录。"],
            ["TC-009", "单据详情", "点击台账明细详情图标。", "弹窗展示来源邮件、来源附件、状态流转日志。"],
            ["TC-010", "压缩包附件", "在详情页展开 ZIP 来源附件。", "展示压缩包内部文件列表，并可切换预览文件内容。"],
        ],
    )

    add_heading(doc, "十、非功能要求", 1)
    add_bullets(
        doc,
        [
            "列表查询响应时间：普通筛选不超过 2 秒，复杂分组查询不超过 5 秒。",
            "导出任务需支持大数据量异步处理，前端展示导出进度或完成提示。",
            "系统日志不可由普通用户编辑或删除。",
            "数据权限必须在后端接口层校验，不能只依赖前端隐藏。",
            "页面在 1366px 及以上分辨率保持字段可读；横向字段较多时使用横向滚动。",
        ],
    )

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    build_doc()
