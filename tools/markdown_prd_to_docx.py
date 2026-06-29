from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(46, 116, 181)
DARK_ACCENT = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
TABLE_FILL = "F2F4F7"
BORDER = "D9E2F3"


def set_run_font(run, size=None, bold=None, color=None, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
      tc_mar = OxmlElement("w:tcMar")
      tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, DARK_ACCENT, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_rich_text(paragraph, text, size=None, bold=None, color=None):
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[1:-1] if is_code else part[2:-2] if is_bold else part
        run = paragraph.add_run(clean)
        set_run_font(
            run,
            size=size,
            bold=True if is_bold else bold,
            color=color,
            name="Consolas" if is_code else "Calibri",
            east_asia="Microsoft YaHei",
        )
        if is_code:
            run.font.size = Pt(size or 10)


def split_table_row(line):
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_table_separator(line):
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def column_widths(headers):
    count = len(headers)
    if count <= 1:
        return [6.5]
    if count == 2:
        return [1.9, 4.6]
    if count == 3:
        return [1.6, 1.7, 3.2]
    if count == 4:
        return [1.2, 1.7, 1.6, 2.0]
    if count == 5:
        return [1.0, 1.25, 1.35, 1.45, 1.45]
    return [6.5 / count] * count


def add_markdown_table(doc, rows):
    headers = split_table_row(rows[0])
    body = [split_table_row(line) for line in rows[2:]]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    widths = column_widths(headers)

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich_text(p, header, size=10, bold=True)

    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(headers)]):
            p = cells[idx].paragraphs[0]
            if len(value) < 12:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, value, size=10)

    set_table_width(table, widths)
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=8, line=1.0)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, size=9.5, name="Consolas", east_asia="Microsoft YaHei", color=RGBColor(31, 41, 55))


def build_docx(md_path: Path, out_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "POS数据管理平台 PRD"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [stripped, lines[i + 1].strip()]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=3, line=1.1)
                add_rich_text(p, text, size=22, bold=True, color=DARK_ACCENT)
                subtitle = doc.add_paragraph()
                set_paragraph_spacing(subtitle, before=0, after=16, line=1.1)
                add_rich_text(subtitle, "由 Markdown PRD 转换生成，用于开发交付与评审。", size=10, color=MUTED)
            else:
                style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
                doc.add_paragraph(text, style=style)
            i += 1
            continue

        checkbox = re.match(r"^-\s+\[\s\]\s+(.+)$", stripped)
        if checkbox:
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, f"☐ {checkbox.group(1)}")
            set_paragraph_spacing(p, after=4, line=1.10)
            i += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, bullet.group(1))
            set_paragraph_spacing(p, after=4, line=1.10)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, numbered.group(1))
            set_paragraph_spacing(p, after=4, line=1.10)
            i += 1
            continue

        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        set_paragraph_spacing(p, before=0, after=6, line=1.10)
        i += 1

    doc.save(out_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: markdown_prd_to_docx.py input.md output.docx")
    build_docx(Path(sys.argv[1]), Path(sys.argv[2]))
