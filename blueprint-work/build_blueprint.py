from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path('/Users/aiden/Desktop/Demo/PosDataSystem')
OUT = ROOT / 'outputs/blueprint/好丽友POS数据管理平台项目一阶段蓝图计划书.docx'
BLUE = '155EEF'
DARK = '1D2129'
MID = '4E5969'
PALE = 'EAF3FF'
LINE = 'B8C8E8'
LIGHT = 'F5F8FC'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for tag, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_repeat_table_header_safe(row):
    set_repeat_table_header(row)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    cant_split.set(qn('w:val'), 'true')
    tr_pr.append(cant_split)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width_cm * 567)))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(int(sum(widths) * 567)))
    tbl_w.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(width * 567)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def set_run_font(run, name='Arial Unicode MS', size=10.5, bold=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_text(paragraph, text, bold=False, color=DARK, size=10.5, align=None):
    paragraph.clear()
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_field(paragraph, instruction, placeholder=''):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = placeholder
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])


def add_top_rule(header):
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(p, '好丽友POS数据管理平台项目｜一阶段蓝图计划书', color=MID, size=8.5)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '5')
    bottom.set(qn('w:color'), DARK)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_section(section, first=False):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.2)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    section.different_first_page_header_footer = first
    if not first:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        add_top_rule(section.header)
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(p, 'PAGE', '1')
        for run in p.runs:
            set_run_font(run, name='Times New Roman', size=9, color=MID)


def apply_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial Unicode MS'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True
    for name, size, color, before, after in [
        ('Heading 1', 16, DARK, 14, 8),
        ('Heading 2', 13, BLUE, 12, 6),
        ('Heading 3', 11, DARK, 9, 4),
    ]:
        style = styles[name]
        style.font.name = 'Arial Unicode MS'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_body(doc, text, indent=True, bold_prefix=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.5
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(3)
        set_paragraph_text(p, item)


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn('w:num'))
    new_num_id = max([int(n.get(qn('w:numId'))) for n in nums] + [0]) + 1
    new_num = OxmlElement('w:num')
    new_num.set(qn('w:numId'), str(new_num_id))
    abstract = OxmlElement('w:abstractNumId')
    abstract.set(qn('w:val'), '7')
    new_num.append(abstract)
    override = OxmlElement('w:lvlOverride')
    override.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:startOverride')
    start.set(qn('w:val'), '1')
    override.append(start)
    new_num.append(override)
    numbering.append(new_num)
    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        num_id = OxmlElement('w:numId')
        num_id.set(qn('w:val'), str(new_num_id))
        num_pr.extend([ilvl, num_id])
        p_pr.append(num_pr)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(3)
        set_paragraph_text(p, item)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, [16.7])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title + '：')
    set_run_font(r, name='Arial Unicode MS', bold=True, color=BLUE)
    r = p.add_run(text)
    set_run_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header_safe(hdr)
    set_row_cant_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, PALE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_paragraph_text(cell.paragraphs[0], text, bold=True, color=BLUE, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ridx, row in enumerate(rows):
        added_row = table.add_row()
        set_row_cant_split(added_row)
        cells = added_row.cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            if ridx % 2 == 1:
                set_cell_shading(cell, LIGHT)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_text(cell.paragraphs[0], value, size=font_size, align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_flow(doc, steps):
    table = doc.add_table(rows=1, cols=len(steps))
    table.style = 'Table Grid'
    table.autofit = False
    widths = [16.7 / len(steps)] * len(steps)
    set_table_width(table, widths)
    for idx, step in enumerate(steps):
        cell = table.cell(0, idx)
        set_cell_shading(cell, BLUE if idx % 2 == 0 else '377CFB')
        set_cell_margins(cell, 160, 70, 160, 70)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        text = f'{idx + 1}\n{step}'
        p = set_paragraph_text(cell.paragraphs[0], text, bold=True, color='FFFFFF', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.line_spacing = 1.15
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_picture(doc, path, caption):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(16.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(cap, caption, color=MID, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.keep_with_next = False


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    apply_styles(doc)
    configure_section(doc.sections[0], first=True)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = ROOT / 'assets/logo.png'
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Cm(4.5))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    set_paragraph_text(p, '好丽友POS数据管理平台项目', bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    set_paragraph_text(p, '一阶段蓝图计划书', bold=True, color=BLUE, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(12)
    p = doc.add_paragraph()
    set_paragraph_text(p, 'Blueprint & Implementation Plan · Phase 1', color=MID, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(16)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(130)
    set_paragraph_text(p, '版本：V1.0（评审稿）', color=MID, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    set_paragraph_text(p, '2026年8月', bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Signed page section
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, first=False)
    h = doc.add_heading('签 署 页', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, '本页用于项目蓝图评审、范围确认及后续实施授权。签署代表相关方已理解本蓝图所描述的一阶段目标、范围、边界和验收原则。')
    rows = [['编制', '', '日期', ''], ['校对', '', '日期', ''], ['审核', '', '日期', ''], ['会签', '', '日期', ''], ['批准', '', '日期', '']]
    table = add_table(doc, ['角色', '签字', '角色', '签字'], rows, [2.6, 5.6, 2.6, 5.9], 10)
    for row in table.rows[1:]:
        row.height = Cm(1.45)
    doc.add_heading('文档控制', level=2)
    add_table(doc, ['项目', '内容'], [
        ['文档名称', '好丽友POS数据管理平台项目一阶段蓝图计划书'],
        ['版本/状态', 'V1.0 / 评审稿'],
        ['适用阶段', '当前整个项目的第一阶段'],
        ['编制依据', '参考蓝图模板、当前项目原型代码、功能描述清单及设计验收记录'],
        ['预期读者', '项目发起人、业务负责人、POS担当、营业担当、产品、开发、测试、实施及运维人员'],
        ['保密级别', '项目内部'],
    ], [4.0, 12.7], 9.5)
    doc.add_heading('版本记录', level=2)
    add_table(doc, ['版本', '日期', '变更说明', '编制人'], [['V1.0', '2026-08-05', '基于当前项目形成一阶段完整蓝图与实施计划', '待填写']], [2.0, 3.2, 8.7, 2.8], 9.2)

    # TOC
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, first=False)
    h = doc.add_heading('目录', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, ['章节', '内容'], [
        ['1', '项目概述'], ['2', '一阶段总体蓝图'], ['3', '数据接入与自动处理蓝图'],
        ['4', '业务应用蓝图方案'], ['5', '数据蓝图'], ['6', '权限、安全与审计蓝图'],
        ['7', '非功能需求蓝图'], ['8', '一阶段实施计划'], ['9', '测试与验收方案'],
        ['10', '风险、依赖与待确认事项'], ['11', '附件'],
    ], [2.0, 14.7], 9.8)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', '右键选择“更新域”以刷新目录')
    add_callout(doc, '说明', '目录与页码为 Word 域。首次打开文档时可按 Ctrl+A 后按 F9 刷新。')

    # Body starts
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, first=False)

    doc.add_heading('1. 项目概述', level=1)
    doc.add_heading('1.1 文档目的', level=2)
    add_body(doc, '本文档在需求调研和当前产品原型的基础上，对好丽友POS数据管理平台项目第一阶段进行蓝图设计，并给出实施、测试、上线和验收计划。文档用于统一业务目标、功能范围、数据口径、处理流程、权限边界和交付标准，是后续详细设计、开发、测试、用户验收及项目管理的共同基线。')
    doc.add_heading('1.2 项目背景', level=2)
    add_body(doc, '当前终端门店POS销售数据主要由业务人员拜访门店、合作方反馈或周期性文件提交获得，源文件以 Excel、CSV、PDF、压缩包等形式为主。不同客户系统的字段名称、表头位置、数据格式、门店编码和产品编码存在差异；在数据进入分析环节前，通常需要经历收件确认、解压拆分、字段识别、门店匹配、产品匹配、质量校验、异常沟通、人工修正、汇总输出等多道处理。')
    add_body(doc, '当月文件数量大、来源分散且问题类型多样时，纯人工处理容易形成周期长、重复劳动多、进度不透明、版本不可追溯和质量口径不一致等问题。项目拟通过标准化平台、自动化流程与AI辅助识别，建立从文件接入到标准POS明细沉淀的闭环。')
    add_callout(doc, '一阶段定位', '搭建可运行、可追踪、可质检、可沉淀的POS数据处理主流程，优先解决“文件收得到、数据看得懂、异常找得到、责任追得清、结果导得出”。')
    doc.add_heading('1.3 建设目标', level=2)
    add_table(doc, ['目标维度', '一阶段目标', '衡量方式'], [
        ['效率', '通过自动解压、文件识别、字段映射、门店匹配和批量校验减少重复操作', '处理时长、自动识别率、人工介入比例'],
        ['质量', '以统一字段和规则识别缺失、格式、重复、映射、数值及逻辑异常', '质检通过率、异常闭环率、抽检准确率'],
        ['透明', '对文件收取、待处理、质检、驳回、复核和同步状态全过程可视', '状态覆盖率、超时记录、操作日志完整率'],
        ['沉淀', '形成标准POS明细台账、门店对照、字段列表和月度提交监控', '标准明细完整率、可导出性、分析口径一致性'],
        ['治理', '建立角色权限、组织数据范围、字段配置和审计日志', '越权测试、审计可追溯、配置生效正确性'],
    ], [2.0, 9.3, 5.4], 9)
    doc.add_heading('1.4 一阶段建设范围', level=2)
    add_table(doc, ['范围类型', '纳入一阶段', '说明'], [
        ['业务功能', '首页概览、文件收取、质量检查、台账与汇总、数据分析、系统设置', '以当前项目原型已覆盖的完整主流程为基线'],
        ['数据对象', '文件、附件、门店、经销商、组织、产品、标准POS明细、异常、处理记录', '统一标识、状态和关联关系'],
        ['自动化能力', '文件/压缩包解析、字段识别、重复识别、规则校验、匹配建议、结果输出', '低置信度和关键异常进入人工复核'],
        ['权限治理', '用户、角色、功能权限、组织数据权限、系统日志、字段配置', '用户权限从角色继承，数据范围按组织控制'],
        ['终端形态', '桌面Web端；中文/韩文界面基础能力', '移动端不作为一阶段主交付'],
    ], [2.2, 7.0, 7.5], 9)
    doc.add_heading('1.5 非一阶段范围', level=2)
    add_bullets(doc, [
        '面向外部门店的独立移动App或小程序；',
        '全集团级数据中台、复杂主数据治理平台或替换现有ERP/CRM系统；',
        '未经确认的外部系统双向实时写回；',
        '预测销量、智能补货、费用结算等高级经营模型；',
        '完全无人值守的自动放行——关键异常和低置信度结果仍需人工复核。',
    ])
    doc.add_heading('1.6 术语和缩略语', level=2)
    add_table(doc, ['术语', '说明'], [
        ['POS', 'Point of Sale，本项目指终端门店销售明细数据。'],
        ['标准POS明细', '经解析、字段标准化、主数据匹配和质量校验后形成的统一粒度销售记录。'],
        ['文件箱/收件箱', '集中展示收到的原始文件、附件及其处理状态的工作区。'],
        ['ACC', '业务侧客户或渠道识别字段，具体编码体系以主数据确认为准。'],
        ['客户交易处', '客户系统侧用于识别门店或交易主体的编码/名称。'],
        ['好丽友交易处', '平台侧标准主数据中的交易主体编码/名称。'],
        ['质检', '对结构、必填、格式、主数据、重复、数值和逻辑一致性进行检查。'],
        ['RPA', 'Robotic Process Automation，用于执行可规则化、重复性的文件和数据处理动作。'],
    ], [3.5, 13.2], 9.2)

    page_break(doc)
    doc.add_heading('2. 一阶段总体蓝图', level=1)
    doc.add_heading('2.1 设计原则', level=2)
    add_table(doc, ['原则', '蓝图要求'], [
        ['标准先行', '先确定标准字段、状态、质量规则和主数据关系，再实施自动化。'],
        ['人机协同', 'AI提供识别、匹配和异常建议；业务人员对低置信度及关键结果负责确认。'],
        ['过程留痕', '原文件、解析结果、修改前后值、操作人、时间、驳回原因和复核结果可追溯。'],
        ['配置驱动', '字段别名、是否必填、是否参与质检和是否显示在台账尽量配置化。'],
        ['权限内生', '功能权限与组织数据权限在所有列表、详情、导出及操作中一致执行。'],
        ['渐进上线', '按代表性区域和文件模板试点，经过回归与并行核对后逐步扩围。'],
    ], [3.0, 13.7], 9.4)
    doc.add_heading('2.2 业务能力架构', level=2)
    add_table(doc, ['层级', '能力模块', '主要能力'], [
        ['用户触点', 'Web工作台 / 中韩文界面', '登录、导航、全局年月、消息提示、智能辅助'],
        ['业务应用', '首页概览 / 文件收取 / 质量检查 / 台账与汇总 / 数据分析', '收取跟踪、处理闭环、质量控制、数据沉淀与监控'],
        ['平台治理', '用户 / 角色权限 / 字段配置 / 系统日志', '身份、授权、组织范围、规则配置、审计追溯'],
        ['智能处理', '解析 / 拆分 / 识别 / 映射 / 校验 / 建议', 'RPA执行固定流程，AI处理非标准结构与辅助判断'],
        ['数据服务', '原始区 / 标准区 / 质量区 / 台账区 / 分析区', '原始留存、标准化、异常记录、结果输出与统计'],
    ], [2.2, 5.6, 8.9], 9)
    doc.add_heading('2.3 端到端业务流程', level=2)
    add_flow(doc, ['文件接入', '预检与解压', '结构解析', '主数据匹配', '质量检查', '人工处置', '复核通过', '台账沉淀'])
    add_body(doc, '主流程从文件接入开始。系统首先对文件类型、完整性、重复性和压缩包可用性进行预检；通过后完成表头、字段和数据区域识别，并结合门店、经销商和产品主数据进行匹配。规则校验和AI辅助判断生成质检结果；异常进入人工处置或驳回流程，通过复核的数据写入标准POS台账并用于分析与导出。')
    doc.add_heading('2.4 角色与职责', level=2)
    add_table(doc, ['角色', '主要职责', '典型权限范围'], [
        ['POS担当', '统筹收取、处理、质检、复核、台账和跨区域问题', '全部或指定组织数据；关键处理与审批权限'],
        ['营业担当', '补充门店/经销商信息，处理被分派或被驳回的问题', '所属Team/区域/营业所及下属数据'],
        ['区域/营业所管理者', '查看进度、质量和提交监控，推动问题闭环', '所辖组织数据的查看与导出'],
        ['系统管理员', '用户、角色、字段配置、日志与基础运维', '系统设置与全局审计权限'],
        ['审计/只读用户', '查看台账、质量记录和操作日志', '受控只读范围'],
    ], [3.0, 7.6, 6.1], 9.2)
    doc.add_heading('2.5 逻辑技术架构', level=2)
    add_table(doc, ['架构层', '组成', '一阶段设计要点'], [
        ['展现层', '桌面Web、响应式页面、中韩文资源', '统一导航、筛选、表格、详情抽屉、对话框与状态提示'],
        ['应用层', '收件、解析、匹配、质检、工作流、台账、分析、权限', '模块解耦，状态驱动，支持批量操作'],
        ['智能/自动化层', 'OCR/表格识别、字段映射、相似匹配、规则引擎、RPA任务', '输出置信度、原因和建议，关键节点可人工接管'],
        ['数据层', '文件存储、业务库、规则配置、日志、分析视图', '原始与标准分层，版本可追溯，逻辑删除/归档'],
        ['集成层', '邮件/上传入口、外部主数据、导出接口（预留）', '接口鉴权、幂等、失败重试和任务监控'],
        ['安全运维层', '身份认证、RBAC、数据权限、审计、监控、备份', '最小权限、敏感操作留痕、异常告警'],
    ], [2.4, 5.3, 9.0], 8.9)

    page_break(doc)
    doc.add_heading('3. 数据接入与自动处理蓝图', level=1)
    doc.add_heading('3.1 数据来源与接入方式', level=2)
    add_table(doc, ['来源', '接入方式', '一阶段处理', '边界/依赖'], [
        ['业务/合作方邮件', '收件箱同步或人工下载后上传', '记录来源、主题、时间、附件与处理状态', '邮件连接方式和账号安全策略待确认'],
        ['人工上传', '平台文件箱上传', '支持单文件、批量文件和压缩包', '单文件大小及批量上限待确认'],
        ['历史存量', '按批次导入', '预检、去重、解析、质检后入账', '需提供样本与历史口径'],
        ['外部系统', '接口/文件交换（预留）', '一阶段完成接口边界和扩展点', '是否实施实时接口需单独确认'],
    ], [3.0, 4.0, 6.0, 3.7], 8.8)
    doc.add_heading('3.2 支持文件与预检', level=2)
    add_bullets(doc, [
        '结构化文件：XLSX、XLS、CSV；根据模板识别表头、数据区、合并单元格和多工作表。',
        '非结构化/半结构化文件：PDF；优先识别可提取表格，扫描件是否启用OCR由样本验证决定。',
        '压缩文件：ZIP；记录解压结果、内部子文件、损坏或加密异常，并支持部分驳回。',
        '预检项：扩展名、MIME类型、大小、空文件、重复文件、文件损坏、加密、宏风险和病毒扫描（部署环境支持时）。',
    ])
    doc.add_heading('3.3 自动化处理链路', level=2)
    add_numbered(doc, [
        '文件登记：生成唯一单据号，保存原始文件、来源、年月、提交人和组织归属。',
        '预检与去重：计算文件指纹，识别重复附件并提供覆盖、忽略或保留新版本的处置选项。',
        '拆包与拆表：解压ZIP，识别工作表和门店边界，必要时拆分为单门店处理单元。',
        '结构解析：识别表头、字段别名、数据起始行、合计行和空行，生成结构化结果。',
        '标准映射：将客户字段映射到标准字段，将客户门店/交易处映射到好丽友标准主数据。',
        '质量校验：执行必填、格式、范围、重复、主数据、一致性和业务逻辑规则。',
        '人机协同：自动通过高置信度且无异常的数据；低置信度或有异常的数据进入人工工作台。',
        '复核与入账：人工修改保留前后值和原因；通过后写入标准POS明细并生成处理日志。',
    ])
    doc.add_heading('3.4 异常及重复文件处置', level=2)
    add_table(doc, ['场景', '系统识别', '可选处置', '留痕要求'], [
        ['同名同内容', '文件名+指纹一致', '忽略、查看历史', '原单据号、重复时间、操作者'],
        ['同名不同内容', '文件名一致、指纹不同', '作为新版本、覆盖（受权限控制）', '版本链与差异说明'],
        ['压缩包损坏/加密', '无法解压或读取', '整体驳回', '错误原因、原文件保留'],
        ['压缩包部分异常', '部分子文件失败', '部分驳回、正常子文件继续', '子文件级状态'],
        ['非POS文件', '内容/结构不符合POS规则', '归档为非POS、删除（受权限控制）', '分类原因和操作日志'],
        ['字段/主数据异常', '缺失、格式错误或无法匹配', '修改、分派、驳回、提交复核', '前后值、处理意见、责任人'],
    ], [3.0, 4.2, 4.8, 4.7], 8.8)
    doc.add_heading('3.5 AI与RPA使用边界', level=2)
    add_callout(doc, '控制原则', 'AI输出必须包含置信度或异常原因，不能在无审计记录的情况下改写原始文件；关键字段、低置信度匹配和跨主数据冲突必须由授权人员复核。')
    add_table(doc, ['能力', '适合自动化', '需人工确认'], [
        ['文件处理', '收取、登记、解压、拆表、格式预检', '加密文件、损坏文件、异常模板'],
        ['字段识别', '已知别名和稳定模板映射', '新模板、歧义表头、合并区域'],
        ['门店匹配', '编码精确匹配、高置信度组合匹配', '一码多店、名称相似但主体不同'],
        ['质量校验', '确定性规则、重复和格式检查', '业务合理性、例外政策'],
        ['结果放行', '规则确认允许的高置信度批次', '关键异常、批量修改、驳回和删除'],
    ], [3.2, 6.7, 6.8], 9)

    page_break(doc)
    doc.add_heading('4. 业务应用蓝图方案', level=1)
    doc.add_heading('4.1 首页概览', level=2)
    add_body(doc, '首页作为管理驾驶舱，按全局年月和用户组织权限展示平台总体规模、文件收取进度与质量检查进度，并提供到台账、文件收取和质量检查模块的快捷入口。')
    add_table(doc, ['功能区', '指标/内容', '交互要求'], [
        ['平台总览', 'POS明细总数、各区域明细总数', '按年月与组织范围联动；可进入台账'],
        ['文件收取', '应收文件、已收文件、异常文件、收取率、区域进度', '点击区域可查看门店/文件明细'],
        ['质量检查', '应检门店、通过门店、异常门店、通过率、区域校验量', '点击区域可查看门店质检明细'],
        ['状态与更新时间', '当前Team、数据状态、最后更新时间', '全局一致、刷新可追踪'],
    ], [3.2, 7.4, 6.1], 9.2)
    doc.add_heading('4.2 文件收取', level=2)
    add_body(doc, '文件收取模块承担原始文件接入、附件预检、解析结果查看、已匹配/未匹配数据处置以及跨角色流转，是第一阶段的业务入口。')
    add_table(doc, ['子功能', '核心内容', '关键操作/规则'], [
        ['文件箱', '来源、主题/文件名、附件、提交年月、收取时间、状态', '上传、预览/下载、重复处理、单据详情、删除'],
        ['单门店已匹配数据', '已识别门店与标准主数据、解析摘要、质检状态', '查看、编辑、发起质检、提交处理结果'],
        ['单门店未匹配数据', '缺失门店编码、名称不一致、交易处未登记等问题', '补充编码、选择标准门店、校验、驳回、提交复核'],
        ['暂存/异常工作区', '解析异常、待补充和待协同记录', '分派责任人、备注、继续处理'],
        ['归档', '非POS文件、已作废或历史版本', '按权限查看、恢复或留存'],
    ], [3.2, 7.2, 6.3], 8.9)
    doc.add_heading('4.2.1 文件收取状态与流转', level=3)
    add_flow(doc, ['待处理', '校验中', '待通过', '质检中', '已同步'])
    add_body(doc, '异常分支包括校验失败、POS担当待处理、营业担当处理中、部分驳回、整体驳回和作废。状态变更必须记录操作者、时间、处理意见和前后状态；同一单据不得在并发操作中被重复放行。')
    doc.add_heading('4.3 质量检查', level=2)
    add_table(doc, ['子功能', '处理对象', '操作与控制'], [
        ['标准POS表', '已完成结构化并进入标准字段口径的数据', '查看、编辑、提交复核、通过、驳回、详情、删除'],
        ['异常数据', '规则或AI识别出的字段/数据问题', '定位异常字段、查看原因与建议、修改、提交复核、通过/驳回'],
        ['批量处理', '同批次、同规则、同责任人的记录', '批量通过/驳回前二次确认；部分失败需明确反馈'],
        ['修改记录', '所有人工修正', '保存原值、新值、原因、操作者与时间'],
    ], [3.2, 7.0, 6.5], 9)
    doc.add_heading('4.3.1 一阶段质量规则', level=3)
    add_table(doc, ['规则类别', '示例', '处理结果'], [
        ['结构完整性', '缺少表头、无数据区、多表结构无法识别', '进入异常；必要时驳回源文件'],
        ['必填检查', '年月、门店、产品、69码、数量、金额等必填字段为空', '定位字段并要求补充'],
        ['格式检查', '年月非法、编码长度错误、数值列含文本', '自动规范化或人工修正'],
        ['主数据检查', '门店/产品/交易处不存在或多重匹配', '选择标准值或提交主数据维护'],
        ['重复检查', '同年月+门店+产品+关键业务键重复', '标记重复，禁止无说明重复入账'],
        ['数值与逻辑', '数量/金额异常、合计关系不一致、销售成本不合理', '按阈值告警并人工确认'],
        ['跨字段一致性', '69码与产品编码对应关系冲突，一码多品', '高优先级异常，必须复核'],
    ], [3.0, 8.2, 5.5], 8.8)
    doc.add_heading('4.4 台账与汇总', level=2)
    add_body(doc, '台账与汇总模块展示已通过质检并进入标准口径的POS明细。用户可按全局年月、组织、经销商、门店、产品和异常状态筛选，查看单据详情，按权限修改或导出。台账表头由启用且配置为“显示在台账”的标准字段决定。')
    add_table(doc, ['能力', '蓝图要求'], [
        ['查询', '支持年月、Team、区域、营业所、经销商、门店编码/名称、产品编码/名称、69码等组合筛选。'],
        ['详情', '回溯原始文件、解析批次、质检记录、修改记录和当前标准值。'],
        ['导出', '导出范围必须与当前筛选和用户数据权限一致；记录导出人、时间和条件。'],
        ['编辑', '仅授权角色可编辑；修改后按配置决定是否重新进入质检。'],
        ['汇总', '按业务确认的维度生成月度汇总；汇总口径与标准明细保持可追溯。'],
    ], [3.4, 13.3], 9.2)
    doc.add_heading('4.5 数据分析', level=2)
    add_table(doc, ['分析页签', '用途', '主要字段/筛选'], [
        ['门店对照表', '查看客户系统门店与好丽友标准交易处的匹配关系', '客户系统、客户门店编码/名称、ACC、标准交易处、匹配状态、操作时间'],
        ['月度门店', '按月份查看门店映射及组织归属快照', '年月、Team、区域、营业所、匹配状态'],
        ['字段列表', '检查各客户/文件在不同月份可提供的字段及备注', '年月、字段、可用性、组合条件'],
        ['门店提交监控', '按月份跟踪门店是否提交POS数据', '月份范围、门店、提交/未提交状态'],
    ], [3.2, 6.2, 7.3], 9)
    add_picture(doc, ROOT / 'implementation-blue-white.png', '图 4-1 当前项目数据分析界面（门店对照表）')
    doc.add_heading('4.6 系统设置', level=2)
    add_table(doc, ['子模块', '功能', '关键控制'], [
        ['用户管理', '新建、编辑、启停、重置密码、删除、权限预览', '用户不单独配置数据范围，默认继承角色'],
        ['角色权限', '配置菜单/动作权限和组织数据范围', '支持全部数据或按Team/区域/营业所选择'],
        ['字段配置', '标准字段名、别名、类型、必填、参与质检、台账显示、启停与排序', '变更后同步影响质检规则与台账表头；需审计'],
        ['系统日志', '按时间、用户、模块、动作和关键词查询并导出', '记录结果、目标、明细、修改前后值'],
    ], [3.2, 7.1, 6.4], 9)
    doc.add_heading('4.7 AI辅助与交互', level=2)
    add_body(doc, '平台可提供全局AI辅助入口，用于解释异常、提示处理步骤、根据上下文定位单据或给出字段匹配建议。一阶段AI不替代业务审批，不直接绕过权限执行删除、放行或跨组织查询；所有实际写操作必须通过明确的业务动作并进入日志。')
    doc.add_heading('4.8 登录与多语言', level=2)
    add_bullets(doc, [
        '支持账号密码登录、密码显示/隐藏、记住账号或凭据策略（需符合安全要求）。',
        '登录后按用户角色和组织数据范围加载导航与数据。',
        '一阶段保留中文/韩文界面切换能力；业务数据本身不强制翻译。',
        '登录失败、账号停用、会话过期和无权限访问应给出明确提示。',
    ])
    add_picture(doc, ROOT / 'implementation-login-redesign.png', '图 4-2 当前项目登录界面')

    page_break(doc)
    doc.add_heading('5. 数据蓝图', level=1)
    doc.add_heading('5.1 数据分层与生命周期', level=2)
    add_flow(doc, ['原始文件区', '解析暂存区', '标准化区', '质量异常区', '标准台账区', '分析服务区'])
    add_body(doc, '原始文件区保存收到的原文件及元数据；解析暂存区保存拆包、拆表和结构化中间结果；标准化区保存字段与主数据映射结果；质量异常区保存规则命中和处置记录；通过质检的数据写入标准台账区，并通过分析视图和导出服务提供给业务。各层之间使用批次号、文件号、单据号和明细行号贯通。')
    doc.add_heading('5.2 核心数据实体', level=2)
    add_table(doc, ['实体', '关键属性', '主要关系'], [
        ['文件批次', '批次号、年月、来源、提交人、组织、状态', '包含一个或多个文件/附件'],
        ['文件/附件', '文件号、文件名、指纹、类型、大小、版本、解析状态', '来源于批次，可拆分为门店处理单元'],
        ['门店处理单元', '单据号、客户门店、标准门店、责任人、流程状态', '关联解析结果、异常和处理记录'],
        ['标准POS明细', '年月、组织、经销商、门店、产品、69码、数量、金额、成本', '来源于已通过质检的解析明细'],
        ['异常记录', '规则、字段、原值、建议、级别、状态', '关联文件、单据或明细行'],
        ['主数据映射', '客户系统/编码/名称、标准编码/名称、有效期、状态', '用于门店和产品匹配'],
        ['操作日志', '用户、模块、动作、目标、结果、前后值、时间', '关联业务对象并支持审计'],
    ], [3.0, 8.2, 5.5], 8.8)
    doc.add_heading('5.3 标准POS字段建议', level=2)
    add_table(doc, ['分类', '字段', '必填建议', '说明'], [
        ['时间', '年月', '是', '统一为 YYYY-MM，决定业务归属周期'],
        ['组织', 'ACC、经销商名称、Team、区域、营业所', '按口径', '组织层级由标准主数据派生'],
        ['门店', '客户门店编码/名称、标准门店编码/名称', '是', '保留客户值与标准值，便于追溯'],
        ['产品', '产品编码、产品名称、69码', '是', '校验编码对应关系和一码多品冲突'],
        ['销售', '销售数量、销售金额、销售成本', '是/按模板', '统一为数值类型，明确单位与小数精度'],
        ['追溯', '批次号、文件号、单据号、源行号、版本', '是', '系统生成，不由业务文件覆盖'],
        ['质量', '质检状态、异常数、置信度、最后处理人/时间', '是', '用于过程控制和查询'],
    ], [2.2, 6.2, 2.6, 5.7], 8.7)
    doc.add_heading('5.4 门店及交易处匹配', level=2)
    add_numbered(doc, [
        '优先使用客户系统+客户门店编码或客户交易处编码进行精确匹配。',
        '编码缺失时，使用客户系统、门店名称、地址、经销商和组织归属进行组合匹配。',
        '组合匹配必须输出候选项和置信度；低于阈值或存在多个候选时禁止自动确认。',
        '人工确认后沉淀为有效期可管理的映射关系，后续同类文件复用。',
        '标准主数据发生合并、停用或更名时，保留历史映射和有效期，避免改写历史台账。',
    ])
    doc.add_heading('5.5 数据状态模型', level=2)
    add_table(doc, ['对象', '主要状态', '状态控制'], [
        ['文件', '待处理、校验中、正常、待处理、部分驳回、整体驳回、作废、已处理', '由预检/解析和人工处置共同驱动'],
        ['门店单据', '待通过、POS担当待处理、营业担当处理中、质检中、已同步、校验失败', '按角色和权限限定可执行动作'],
        ['异常', '待处理、处理中、待复核、已通过、已驳回、已关闭', '关闭前必须有处理结果或豁免理由'],
        ['台账明细', '有效、待复核、已更正、已作废', '更正生成版本链，禁止无痕覆盖'],
    ], [3.0, 8.0, 5.7], 9)
    doc.add_heading('5.6 数据保留与归档原则', level=2)
    add_bullets(doc, [
        '原始文件、标准明细、异常记录和审计日志的保留期限由好丽友数据管理和合规要求确定。',
        '业务删除默认采用逻辑删除或作废；物理删除仅限运维和合规批准场景。',
        '导出文件应设置有效期和访问控制，避免长期散落。',
        '历史批次归档后仍可按权限查询，恢复操作需记录日志。',
    ])

    page_break(doc)
    doc.add_heading('6. 权限、安全与审计蓝图', level=1)
    doc.add_heading('6.1 权限模型', level=2)
    add_body(doc, '平台采用“用户—角色—功能权限—组织数据范围”的RBAC模型。用户分配一个或多个经确认的业务角色，功能权限细化到模块/页面/动作，数据权限按Team、区域、营业所等组织范围控制。所有列表、详情、统计、批量操作和导出必须使用同一数据权限过滤规则。')
    add_table(doc, ['权限层', '控制内容', '示例'], [
        ['菜单/页面', '是否可见和可访问', '质量检查、系统设置'],
        ['业务动作', '查看、新建、编辑、质检、通过、驳回、删除、导出', '营业担当可查看但不一定可通过'],
        ['数据范围', '全部数据或指定Team/区域/营业所', '华北Team、石家庄营业所'],
        ['字段/敏感操作', '字段可见、批量导出、覆盖、删除、角色配置', '仅管理员或授权角色'],
    ], [3.1, 7.3, 6.3], 9)
    doc.add_heading('6.2 安全要求', level=2)
    add_bullets(doc, [
        '身份认证：密码复杂度、失败锁定、会话超时、停用账号即时失效；是否接入企业SSO待确认。',
        '传输与存储：生产环境使用HTTPS；数据库凭据、邮箱凭据和接口密钥集中管理，不写入前端。',
        '文件安全：限制扩展名与大小，检测损坏、宏和恶意内容；原始文件下载受权限控制。',
        '数据隔离：服务端实施组织数据权限，禁止仅依赖前端隐藏。',
        '敏感操作：批量通过、驳回、覆盖、删除、角色修改和导出需要二次确认与审计。',
        '备份恢复：数据库与文件存储制定备份、恢复和恢复演练机制。',
    ])
    doc.add_heading('6.3 审计日志', level=2)
    add_table(doc, ['日志类型', '记录内容', '检索维度'], [
        ['登录登出', '账号、时间、结果、来源IP/设备（部署支持时）', '日期、用户、结果'],
        ['业务操作', '模块、动作、目标、结果、备注', '用户、模块、动作、单据号'],
        ['数据修改', '字段、修改前值、修改后值、原因', '文件/门店/明细、操作者'],
        ['审批流转', '前后状态、处理意见、责任人', '批次、单据、状态'],
        ['配置变更', '角色、权限、字段规则和系统配置前后值', '配置项、管理员、日期'],
        ['导出', '导出人、条件、范围、文件名和结果', '用户、模块、日期'],
    ], [3.2, 8.2, 5.3], 9)

    page_break(doc)
    doc.add_heading('7. 非功能需求蓝图', level=1)
    add_table(doc, ['类别', '一阶段基线要求', '验证方式'], [
        ['性能', '常用列表查询、筛选和详情在正常网络与目标数据量下保持可接受响应；大批量任务异步执行并可查看进度', '性能测试、任务并发测试；具体指标待容量确认'],
        ['可用性', '关键任务失败可重试；页面刷新不造成重复提交；状态和错误提示清晰', '异常注入、幂等测试、用户体验走查'],
        ['可扩展性', '字段、别名、规则、组织范围和文件模板尽量配置化；预留外部接口扩展点', '新增模板/字段验证'],
        ['可维护性', '模块化日志、任务追踪、配置变更审计、环境参数分离', '运维手册和故障演练'],
        ['兼容性', '支持企业标准桌面浏览器；Excel/PDF/ZIP以样本库验证', '浏览器矩阵与文件样本回归'],
        ['可访问性', '表单标签、键盘操作、焦点状态、颜色对比和明确错误信息', '键盘与可访问性检查'],
        ['国际化', '中文/韩文资源分离，界面切换不影响数据与权限', '双语页面回归'],
        ['数据质量', '标准字段、映射、规则和异常原因可解释、可追溯', '抽样对账与规则命中核验'],
    ], [2.5, 9.1, 5.1], 8.7)
    add_callout(doc, '容量假设', '参考资料提到月度文件量可能达到4,000+。该数字需由业务再次确认，并据此确定单文件大小、并发上传、月度明细量、历史保留年限和性能验收指标。')

    page_break(doc)
    doc.add_heading('8. 一阶段实施计划', level=1)
    doc.add_heading('8.1 实施策略', level=2)
    add_body(doc, '建议采用“蓝图确认—详细设计—核心流程开发—试点验证—并行核对—分批上线”的方式实施。计划以相对周次W1–W12表达，正式启动日期确认后再转换为日历计划；如外部接口、主数据清洗或文件样本复杂度超出假设，应通过变更流程调整。')
    doc.add_heading('8.2 阶段计划（建议基线）', level=2)
    add_table(doc, ['阶段', '周次', '主要工作', '里程碑/输出'], [
        ['1. 蓝图确认', 'W1–W2', '业务访谈、样本盘点、范围/流程/字段/规则确认、原型评审', '蓝图签署、待确认项关闭清单'],
        ['2. 详细设计', 'W3', '数据模型、接口、任务、权限、日志、异常和页面详细设计', '详细设计与测试策略评审'],
        ['3. 核心开发', 'W4–W7', '文件接入、解析匹配、质量检查、工作流、台账、权限', '核心流程联调版本'],
        ['4. 分析与治理', 'W6–W8', '首页、数据分析、字段配置、系统日志、导出、多语言', '功能完整版本'],
        ['5. 测试整改', 'W8–W10', '单元/集成/系统/安全/性能测试，样本回归和缺陷整改', 'SIT通过、UAT候选版本'],
        ['6. UAT与试点', 'W10–W11', '代表区域和模板试点、业务并行核对、培训、上线演练', 'UAT签署、上线审批'],
        ['7. 上线与护航', 'W12', '分批上线、数据初始化、监控、问题响应和复盘', '生产验收、运维移交'],
    ], [2.6, 2.0, 7.8, 4.3], 8.6)
    doc.add_heading('8.3 关键里程碑', level=2)
    add_table(doc, ['里程碑', '完成标准', '建议责任方'], [
        ['M1 蓝图冻结', '范围、流程、字段、角色、外部依赖和验收原则经业务签署', '业务负责人/项目经理'],
        ['M2 核心链路贯通', '代表性文件从接入到台账全流程可运行并可追溯', '产品/开发/测试'],
        ['M3 SIT通过', '阻断性缺陷清零，主要场景和样本回归通过', '测试负责人'],
        ['M4 UAT通过', '业务关键用户完成场景验证与数据抽检并签署', '业务负责人/关键用户'],
        ['M5 生产验收', '上线稳定期完成、问题闭环、文档和运维移交完成', '项目委员会/运维'],
    ], [3.3, 9.1, 4.3], 9)
    doc.add_heading('8.4 主要交付物', level=2)
    add_bullets(doc, [
        '项目一阶段蓝图计划书及评审纪要；',
        '业务需求/功能清单、字段字典、质量规则清单、状态流转与权限矩阵；',
        '详细设计说明、数据模型、接口说明和部署架构；',
        '可部署的软件版本、数据库脚本、配置和初始化工具；',
        '测试计划、测试用例、测试报告、性能/安全验证记录；',
        'UAT记录、上线方案、回退方案、培训材料、用户手册和运维手册；',
        '问题清单、已知限制、上线后支持与移交记录。',
    ])
    doc.add_heading('8.5 项目职责建议（RACI）', level=2)
    add_table(doc, ['工作项', '业务负责人', 'POS担当/营业担当', '项目团队', '运维/安全'], [
        ['范围与口径确认', 'A', 'R/C', 'R/C', 'I'],
        ['样本与主数据提供', 'A', 'R', 'C', 'I'],
        ['设计、开发与测试', 'C', 'C', 'A/R', 'C'],
        ['UAT与数据抽检', 'A', 'R', 'C', 'I'],
        ['生产部署与安全', 'I', 'I', 'R/C', 'A/R'],
        ['上线推广与问题闭环', 'A', 'R', 'R', 'C'],
    ], [5.0, 3.0, 3.5, 2.9, 2.3], 8.8)
    add_body(doc, '注：R=执行负责，A=最终负责，C=协作/咨询，I=知会。正式项目组织确定后应替换为具体部门和人员。', indent=False)
    doc.add_heading('8.6 上线与迁移策略', level=2)
    add_numbered(doc, [
        '选择文件模板和异常类型具有代表性的区域作为试点，建立样本基线。',
        '准备用户、角色、组织、标准字段、门店/产品映射和规则等初始化数据。',
        '对需要迁移的历史文件分批执行预检、解析、质检和抽样对账，不直接绕过质量流程。',
        '试点期与原人工流程并行核对至少一个业务周期或经业务批准的周期。',
        '达到UAT与数据抽检标准后分批扩围；上线窗口准备回退方案、联系人和问题分级机制。',
        '上线后设置护航期，跟踪解析成功率、异常量、人工介入、处理时长和用户问题。',
    ])

    page_break(doc)
    doc.add_heading('9. 测试与验收方案', level=1)
    doc.add_heading('9.1 测试范围', level=2)
    add_table(doc, ['测试类型', '重点'], [
        ['功能测试', '各模块、状态、操作、筛选、详情、批量、导出、多语言'],
        ['文件样本测试', '不同客户模板、工作表、表头位置、PDF、ZIP、损坏/重复/加密文件'],
        ['数据质量测试', '字段映射、门店/产品匹配、规则命中、修改前后值、入账结果'],
        ['流程与权限测试', '角色动作、组织范围、驳回/复核、并发与幂等、越权访问'],
        ['接口与任务测试', '鉴权、超时、失败重试、重复消息、任务状态和告警'],
        ['性能与稳定性测试', '批量上传、大文件、月度峰值、列表查询、导出和长任务'],
        ['安全与审计测试', '登录、会话、文件安全、敏感操作、日志完整性、凭据保护'],
        ['UAT', '真实业务场景、代表性区域、抽样对账、可操作性和培训效果'],
    ], [3.5, 13.2], 9.2)
    doc.add_heading('9.2 建议验收标准', level=2)
    add_table(doc, ['验收域', '建议通过条件'], [
        ['范围', '本蓝图一阶段范围内的关键功能可用，阻断主流程的缺陷为0。'],
        ['主流程', '代表性文件可从接入、解析、匹配、质检、处置、复核到台账完整贯通。'],
        ['数据', '抽样记录与人工基准对账一致；差异均有解释、规则或已批准豁免。'],
        ['质量', '约定样本库的解析、匹配和规则结果达到双方确认的目标；具体百分比在蓝图冻结前补齐。'],
        ['权限', '角色和组织范围测试通过，无跨组织越权查看、编辑或导出。'],
        ['审计', '关键操作、数据修改、审批、配置和导出日志完整可检索。'],
        ['性能', '在确认容量和并发下达到响应及批处理窗口目标，无不可恢复失败。'],
        ['上线准备', '培训、用户/运维手册、部署与回退方案、备份恢复和支持机制齐备。'],
    ], [3.5, 13.2], 9.1)
    doc.add_heading('9.3 缺陷准出原则', level=2)
    add_bullets(doc, [
        'P0/P1（系统不可用、数据丢失、严重越权、主流程阻断）必须清零。',
        'P2缺陷需完成修复或由业务负责人书面接受并明确修复计划。',
        '展示性轻微问题不得影响关键字段可读性、操作完成和验收场景。',
        '所有遗留问题进入上线清单，明确责任人、影响、临时方案和完成日期。',
    ])

    page_break(doc)
    doc.add_heading('10. 风险、依赖与待确认事项', level=1)
    doc.add_heading('10.1 主要风险及应对', level=2)
    add_table(doc, ['风险', '影响', '应对措施'], [
        ['文件模板数量和变化频繁', '解析准确率下降、开发反复', '建立样本库和模板版本；字段别名与规则配置化；上线后监控新模板'],
        ['主数据不完整或编码不一致', '门店/产品匹配失败', '提前清洗；明确主数据责任人；低置信度人工确认并沉淀映射'],
        ['业务口径未冻结', '质量规则和汇总结果争议', '蓝图阶段形成字段字典、规则清单和签署机制'],
        ['AI误识别或过度自动化', '错误数据进入台账', '置信度门槛、人机复核、抽样检查、禁止AI绕过审批'],
        ['外部接口或邮箱接入延期', '自动收取无法按期上线', '保留人工上传兜底；接口独立联调和模拟器'],
        ['峰值文件与明细量不明确', '性能不足或成本失控', '确认容量模型，批处理异步化，开展峰值压测'],
        ['权限配置复杂', '越权或业务无法操作', '角色模板、权限预览、自动化权限回归和审计'],
        ['用户习惯改变', '上线后绕开系统', '试点、并行期、培训、快速支持和指标跟踪'],
    ], [4.0, 4.7, 8.0], 8.6)
    doc.add_heading('10.2 外部依赖', level=2)
    add_bullets(doc, [
        '业务提供代表性文件样本、异常样本、历史数据量和处理规则。',
        '主数据责任方提供门店、交易处、经销商、产品和组织数据及更新机制。',
        '信息安全/运维确认部署环境、账号体系、邮件接入、文件扫描、备份和日志要求。',
        '业务负责人确认角色、组织范围、驳回/通过权限和验收抽样方法。',
        '若需要外部系统接口，系统所有方提供接口文档、测试环境和联调窗口。',
    ])
    doc.add_heading('10.3 蓝图冻结前待确认事项', level=2)
    add_table(doc, ['编号', '待确认事项', '当前采用的工作假设', '建议责任方'], [
        ['Q1', '一阶段正式启动与目标上线日期', '先按W1–W12相对周次计划', '项目负责人'],
        ['Q2', '真实月度文件数、明细量、峰值并发、单文件大小、历史年限', '参考资料存在“月别文件4,000+”描述，需核实', '业务/运维'],
        ['Q3', '正式数据接入渠道', '人工上传为必备兜底，邮件/接口按确认实施', '业务/IT'],
        ['Q4', '一阶段必须支持的文件格式与样本覆盖', 'XLSX/XLS/CSV/PDF/ZIP；OCR按样本评估', 'POS担当'],
        ['Q5', '标准字段、必填性、单位、小数精度和唯一业务键', '以当前原型字段为起点，不作为最终口径', '业务/数据负责人'],
        ['Q6', '门店/产品/交易处主数据来源及责任人', '由好丽友标准主数据提供并定期更新', '主数据责任方'],
        ['Q7', '自动通过的置信度、关键异常和人工复核边界', '低置信度与主数据冲突必须人工复核', '业务负责人'],
        ['Q8', '汇总报表口径和目标输出模板', '一阶段先确保标准明细导出，汇总模板待确认', 'POS担当'],
        ['Q9', '账号体系、SSO、密码和会话策略', '先支持平台账号，预留SSO', '信息安全'],
        ['Q10', '数据/文件/日志保留期限及合规要求', '不做物理删除，期限待制度确认', '合规/运维'],
        ['Q11', '中文/韩文一阶段覆盖程度', '界面基础双语，业务数据不自动翻译', '业务负责人'],
        ['Q12', '试点区域、关键用户和UAT抽样比例', '选择模板和异常类型具代表性的区域', '项目负责人'],
    ], [1.3, 5.2, 7.1, 3.1], 8.2)
    add_callout(doc, '评审建议', 'Q1–Q12应在蓝图冻结前逐项关闭；无法立即关闭的事项需明确工作假设、影响、责任人和最晚确认日期，并纳入项目风险台账。')

    doc.add_heading('11. 附件', level=1)
    doc.add_heading('11.1 一阶段功能范围矩阵', level=2)
    add_table(doc, ['模块', '子功能', '一阶段优先级', '备注'], [
        ['首页概览', '平台/文件收取/质量检查指标与区域明细', 'Must', '按年月和权限联动'],
        ['文件收取', '文件箱、上传、预览下载、重复处理、详情、删除', 'Must', '原文件留存与状态跟踪'],
        ['文件收取', '已匹配/未匹配门店处理、校验、驳回、复核', 'Must', '端到端主流程'],
        ['质量检查', '标准POS表、异常数据、编辑、通过/驳回、详情', 'Must', '修改留痕'],
        ['台账与汇总', '标准POS明细查询、详情、编辑、导出', 'Must', '导出受权限控制'],
        ['数据分析', '门店对照表、月度门店、字段列表、提交监控', 'Should', '一阶段纳入当前原型能力'],
        ['系统设置', '用户、角色权限、字段配置、系统日志', 'Must', '字段配置入口按权限开放'],
        ['平台能力', '全局年月、中韩文、AI辅助、统一筛选/表格', 'Should', 'AI仅作辅助'],
        ['外部集成', '邮件/主数据/导出接口', 'Conditional', '以Q3、Q6确认结果为准'],
    ], [3.0, 7.2, 2.7, 3.8], 8.7)
    doc.add_heading('11.2 状态与责任方对照', level=2)
    add_table(doc, ['状态', '含义', '主要责任方', '允许的后续动作'], [
        ['待处理', '文件已接入，尚未完成预检/解析或需人工决定', '系统/POS担当', '开始校验、忽略重复、驳回'],
        ['校验中', '系统正在解析、匹配和执行规则', '系统', '查看进度；禁止重复提交'],
        ['待通过', '自动处理完成，等待授权人员确认', 'POS担当', '质检、修改、通过、驳回'],
        ['营业担当处理中', '问题已分派给所属营业团队', '营业担当', '补充/修改、提交复核'],
        ['质检中/待复核', '处理结果等待复核', 'POS担当/复核人', '通过、驳回'],
        ['已同步', '已通过并写入标准台账', '系统', '查看、导出；更正需新流程'],
        ['校验失败/驳回', '文件或数据不符合要求', 'POS担当/营业担当', '查看原因、修复或重新提交'],
        ['作废', '文件或记录不再参与当前业务处理', '授权人员', '按权限查看/恢复；保留日志'],
    ], [3.1, 6.8, 3.0, 3.8], 8.7)
    doc.add_heading('11.3 蓝图结论', level=2)
    add_body(doc, '本蓝图将当前项目第一阶段定义为POS数据处理主流程和基础治理能力的建设阶段。阶段成果不是单一页面集合，而是从文件接入、自动解析、主数据匹配、质量控制、跨角色协同到标准台账与分析监控的完整闭环。项目进入详细设计和开发前，应以本文件的范围、数据口径、状态、权限、待确认事项和验收原则为评审基线。')

    # Update fields on open.
    settings = doc.settings._element
    update = settings.find(qn('w:updateFields'))
    if update is None:
        update = OxmlElement('w:updateFields')
        settings.append(update)
    update.set(qn('w:val'), 'true')

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
