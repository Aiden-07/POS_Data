from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).resolve().parent / "交付文档" / "POS数据管理平台-单据详情统一逻辑及模块显示规则.docx"

BLUE = "2F6BFF"
DARK = "172B4D"
MID = "44546A"
LIGHT = "E8EEF8"
PALE = "F5F7FA"
BORDER = "CDD5DF"
WHITE = "FFFFFF"
GREEN = "16865C"
AMBER = "9A6700"
RED = "C9372C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=10.5, bold=False, color=DARK, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=10.5, bold=False, color=DARK,
             before=0, after=6, line=1.25, align=None, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run_font(p.add_run(text), size=10.2, color=DARK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run_font(p.add_run(text), size=10.2, color=DARK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 11.5}[level], bold=True,
                 color=BLUE if level < 3 else DARK)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        set_cell_shading(header.cells[i], LIGHT)
        p = header.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=9.3, bold=True, color=DARK)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "FAFBFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_run_font(p.add_run(str(value)), size=font_size, color=DARK)
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)
    return table


def add_callout(doc, label, text, tone="blue"):
    fills = {"blue": "EDF4FF", "amber": "FFF7E6", "red": "FFF0F0", "green": "EAF7F2"}
    colors = {"blue": BLUE, "amber": AMBER, "red": RED, "green": GREEN}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fills[tone], size="1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    set_run_font(p.add_run(f"{label}  "), size=10, bold=True, color=colors[tone])
    set_run_font(p.add_run(text), size=10, color=DARK)
    add_para(doc, "", after=3)


def add_field_block(doc, title, common_fields, conditional="无"):
    add_heading(doc, title, 3)
    add_para(doc, "固定显示字段：" + "、".join(common_fields) + "。", size=10.2, after=4)
    add_para(doc, "条件显示字段：" + conditional + "。", size=10.2, after=6)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, before, after in ((1, 16, 18, 10), (2, 13, 14, 7), (3, 11.5, 10, 5)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(p.add_run("POS 数据管理平台 · 单据详情规格  |  "), size=8.5, color="758195")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def module_section(doc, number, name, meta_rows, field_fixed, field_conditional, log_rules, interactions, acceptance):
    add_heading(doc, f"{number}. {name}", 1)
    add_table(doc, ["规则项", "模块规则"], meta_rows, [1900, 7460], font_size=9.4)
    add_field_block(doc, "显示字段", field_fixed, field_conditional)
    add_heading(doc, "日志与流转", 3)
    for item in log_rules:
        add_bullet(doc, item)
    add_heading(doc, "交互规则", 3)
    for item in interactions:
        add_bullet(doc, item)
    add_heading(doc, "模块验收重点", 3)
    for item in acceptance:
        add_bullet(doc, item)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    add_footer(section)

    # Memo masthead
    add_para(doc, "产品 / 研发实现规格", size=10.5, bold=True, color=BLUE, after=12)
    add_para(doc, "POS 数据管理平台", size=14, bold=True, color=MID, after=3)
    add_para(doc, "单据详情统一逻辑及各模块显示规则", size=24, bold=True, color=DARK, after=8, line=1.05)
    add_para(doc, "覆盖文件收取、质量检查、台账与汇总模块", size=12, color=MID, after=16)
    add_table(doc, ["文档信息", "内容"], [
        ["文档用途", "产品评审、UI/交互确认、前后端开发、联调与验收"],
        ["版本", "V1.0（基于 2026-07-21 当前原型与代码逻辑整理）"],
        ["适用范围", "文件箱、单门店已匹配数据、单门店未匹配数据、标准POS表、异常数据、标准POS明细，以及归档（非POS表）"],
        ["实现基线", "统一右侧抽屉；统一原始文件、拆分结果、模块信息、日志流转；按模块注入差异字段与日志"],
    ], [1900, 7460], font_size=9.5)
    add_callout(doc, "结论", "所有模块应复用同一详情容器和数据契约；差异只体现在权限节点、状态、业务字段、日志生成规则及可预览对象。", "blue")

    add_heading(doc, "1. 文档范围与对象模型", 1)
    add_para(doc, "“单据详情”用于追溯一条业务数据从原始文件进入系统，到拆分、匹配、质检、复核和入账的全过程。详情不是编辑页面，不承载状态变更操作；它提供来源证据、当前业务信息和流转日志。")
    add_table(doc, ["对象层级", "定义", "典型入口"], [
        ["父单据", "一次邮件/上传任务，可能包含多个附件", "文件箱父行"],
        ["附件单据", "父单据下的一个 Excel/CSV/ZIP 文件，可存在版本", "文件箱附件行"],
        ["门店单据", "附件拆分后的一家门店数据", "已匹配、未匹配、标准POS表门店视图"],
        ["产品明细", "门店单据中的一条产品销售记录", "标准POS表产品视图、异常产品视图、标准POS明细"],
    ], [1500, 4180, 3680], font_size=9.1)

    add_heading(doc, "2. 统一显示规则", 1)
    add_heading(doc, "2.1 入口与权限", 2)
    for item in [
        "操作列使用“单据详情”图标按钮；无“单据详情”权限时按钮不显示。",
        "点击后必须再次校验“查看”与“单据详情”两个权限；任一缺失则不打开抽屉，并提示“当前账号无单据详情权限”。",
        "数据范围沿用列表查询范围。详情接口不得绕过组织范围、角色范围或记录级授权。",
        "归档（非POS表）属于已匹配数据的特殊只读状态；如后续独立成权限节点，应补充“查看/单据详情”双权限。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 抽屉结构与展示顺序", 2)
    for step in [
        "标题固定为“单据详情”，桌面端宽度建议 680px；内容纵向滚动，关闭后回到原列表位置。",
        "原始文件：始终位于第一段，展示文件名、文件类型与预览入口。",
        "拆分结果：仅多门店文件且至少拆分出 2 条门店记录时展示；支持版本切换。",
        "当前模块信息：展示模块注入的业务字段，字段顺序严格按本规格执行。",
        "当前模块日志：按时间正序展示节点、状态、动作、时间、操作人。",
    ]:
        add_number(doc, step)
    add_callout(doc, "当前实现差异", "现有调用方已构造 moduleFields，但通用 openDocumentDetail 容器尚未渲染“当前模块信息”区。研发实现时必须补齐，否则各模块字段规则不会在页面生效。", "amber")

    add_heading(doc, "2.3 原始文件与预览", 2)
    add_table(doc, ["场景", "取值与行为"], [
        ["显式文件名", "优先 context.sourceFileName，其次 row.sourceFileName / originalFileName / fileName / currentFileName"],
        ["关联文件箱", "通过 sourceEmailId + sourceAttIdx 精确查找；缺失时使用文件名/门店名与主题、正文、附件名做模糊关联"],
        ["仍无法关联", "按合作方、门店、编码、期间生成可读的 .xlsx 兜底名称；字段缺失显示“-”"],
        ["正常附件", "点击打开文件预览；ZIP 显示包内文件并支持逐项预览"],
        ["驳回附件", "禁止预览，提示“文件异常、无法预览”"],
        ["兜底文件", "点击打开当前门店/产品数据预览，不假装存在真实附件下载地址"],
    ], [2150, 7210], font_size=9.2)

    add_heading(doc, "2.4 拆分结果与版本", 2)
    for item in [
        "仅对多门店 Excel 展示；按原始附件分组，版本按 V1、V2…升序排列，默认选中“当前版本”。",
        "版本标签展示“首次上传/驳回重传”与“历史版本/当前版本”；展示上传时间、拆分数量及各状态数量。",
        "每条拆分记录展示单门店 Excel 文件名、门店名、状态、原因；可点击预览来源、已匹配门店或未匹配暂存数据。",
        "状态色：成功类绿色；未匹配/重复/待处理黄色；驳回/失败红色；处理中信息类蓝色。",
        "校验中或驳回且尚无拆分数据时不生成虚假拆分行。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.5 业务字段通用规则", 2)
    add_table(doc, ["规则", "要求"], [
        ["字段顺序", "按模块定义的固定顺序渲染，不随列表自定义列变化"],
        ["空值", "null、undefined、空字符串统一显示“-”；数值 0 必须显示为 0，不得误判为空"],
        ["金额", "详情内统一使用人民币符号与千分位，保留业务要求的小数位；不得出现“￥-”"],
        ["编码", "保留前导零；按字符串处理；长编码允许复制并完整展示 tooltip"],
        ["字段分组", "建议按“基本信息 / 组织与门店 / 产品与金额 / 处理信息”分组；窄屏单列，桌面双列"],
        ["安全", "所有服务端或用户输入必须转义；日志、文件名、备注不得直接拼接未转义 HTML"],
    ], [1750, 7610], font_size=9.2)

    add_heading(doc, "2.6 日志统一规则", 2)
    add_table(doc, ["日志字段", "显示规则"], [
        ["节点", "业务环节名称，例如文件收取、文件校验、门店匹配、质检、提交复核、台账入库"],
        ["状态", "使用模块展示状态；状态色由异常/待办/处理中/成功四类映射"],
        ["动作", "说明发生了什么、为何发生；优先显示真实原因或处理说明，缺失时使用模块兜底文案"],
        ["时间", "优先使用节点真实时间；无真实时间显示“-”，禁止用当前时间伪造"],
        ["操作人", "人工操作显示姓名/角色；系统动作显示“系统”；未知显示“-”"],
        ["排序", "按业务发生顺序正序；历史版本日志保留，不被新版本覆盖"],
    ], [1750, 7610], font_size=9.2)

    add_heading(doc, "3. 权限与模块入口矩阵", 1)
    add_table(doc, ["一级模块", "二级模块", "权限校验", "详情粒度"], [
        ["文件收取", "文件箱", "文件箱:查看 + 单据详情", "父单据 / 附件"],
        ["文件收取", "单门店已匹配数据", "已匹配数据:查看 + 单据详情", "门店"],
        ["文件收取", "单门店未匹配数据", "未匹配数据:查看 + 单据详情", "门店"],
        ["质量检查", "标准POS表", "标准POS表:查看 + 单据详情", "门店 / 产品"],
        ["质量检查", "异常数据", "异常数据:查看 + 单据详情", "门店 / 产品"],
        ["台账与汇总", "标准POS明细", "标准POS明细:查看 + 单据详情", "产品明细"],
        ["文件收取", "归档（非POS表）", "当前复用归档列表可见性；建议独立双权限", "门店 / 文件"],
    ], [1500, 2450, 3000, 2410], font_size=8.8)

    module_section(doc, "4", "文件收取—文件箱", [
        ["入口", "父单据行或附件行的“单据详情”按钮"],
        ["粒度", "父单据：一次邮件/上传；附件：指定 attachmentIndex 的文件"],
        ["状态来源", "父单据 statusText；附件优先 attachment.status"],
        ["原始文件", "父单据展示全部附件；附件详情只聚焦当前附件及其同源版本"],
        ["拆分结果", "父单据可展示所有多门店附件；附件详情只展示当前同源附件版本"],
    ], [
        "单据标题", "上传用户", "年月", "附件数", "上传时间"
    ], "附件粒度追加：附件名称、文件版本、异常说明；附件上传时间优先 uploadedAt，缺失取父单据 provideTime",
    [
        "父单据：文件收取（已接收）→ 文件校验（当前状态）。",
        "普通附件：文件收取 → 文件校验；校验正常时追加门店匹配/拆分结果。",
        "历史版本：文件收取 → 文件处理（原状态/驳回原因）→ 版本切换（历史版本）。",
        "驳回重传：记录旧版本、重传接收、版本切换、重新校验；多门店时补充拆分与重复处理。",
        "重复待处理：保留首次驳回、重新上传、重复识别和覆盖/忽略待办日志。",
    ], [
        "附件正常时点击文件名预览；驳回附件禁止预览。",
        "ZIP 可展开包内 Excel/CSV 并逐项预览。",
        "版本切换只更新拆分面板，不关闭抽屉、不改变列表筛选。",
    ], [
        "父单据与附件详情不会混淆附件范围。",
        "历史版本与当前版本均可追溯，且默认定位当前版本。",
        "异常原因、拆分数量和拆分状态与列表/数据源一致。",
    ])

    module_section(doc, "5", "文件收取—单门店已匹配数据", [
        ["入口", "操作列“单据详情”；归档模式复用同一入口"],
        ["粒度", "门店单据"],
        ["状态来源", "正常模式使用角色可见状态；归档模式固定“已归档”"],
        ["来源关联", "优先 sourceEmailId/sourceAttIdx，否则按文件名或门店名关联"],
    ], [
        "时间", "客户系统", "原始文件名", "当前文件名", "客户门店号", "客户交易处编码", "客户交易处名称", "营业Team", "区域", "营业所", "ACC", "好丽友交易处编码", "好丽友交易处名称", "当前责任人", "最近操作人", "判断说明"
    ], "归档模式：当前责任人取归档处理人，最近操作人显示“-”",
    [
        "已匹配数据：门店匹配完成（待质检）→ 质检中/已质检 → 质检结果已同步至标准POS表或异常表。",
        "尚未质检时日志必须明确“等待质量检查”；已同步时显示实际 qualityRoute/logRoute。",
        "归档：文件收取（如有关联）→ 文件校验（已归档），不追加门店匹配和质检节点。",
    ], [
        "详情只读；文件预览从原始来源打开。",
        "列表处于归档模式时不得错误要求已匹配数据权限；建议统一改为归档专属权限。",
    ], [
        "原始/当前文件名均可追溯，组织字段与门店映射结果一致。",
        "质检路由显示准确，不把异常表数据描述成标准POS表。",
        "归档详情不展示无意义的质检成功日志。",
    ])

    module_section(doc, "6", "文件收取—单门店未匹配数据", [
        ["入口", "暂存/未匹配列表操作列“单据详情”"],
        ["粒度", "门店单据"],
        ["状态来源", "getStashDisplayStatus：待处理 / 处理中 / 待复核等角色展示状态"],
        ["责任人", "由暂存工作流和当前角色计算 currentOwnerName / lastOperatorName"],
    ], [
        "时间", "客户系统", "原始文件名", "当前文件名", "客户门店号", "客户交易处编码", "客户交易处名称", "营业Team", "区域", "营业所", "ACC", "好丽友交易处编码", "好丽友交易处名称", "异常说明", "当前责任人", "最近操作人"
    ], "工作流存在驳回时，日志展示 rejectNote；营业提交复核时展示 salesSubmitNote",
    [
        "未匹配数据生成：初始待处理，动作说明门店编码缺失或未匹配主数据。",
        "被 POS 驳回：追加“POS驳回/处理中”，展示真实驳回原因、时间和操作人。",
        "营业提交：追加“提交复核/待复核”，展示处理说明、时间和提交人。",
    ], [
        "详情不直接提供编辑、校验、驳回按钮；这些动作仍在列表操作区完成。",
        "原始文件/拆分行可预览对应来源；不可关联时使用兜底预览。",
    ], [
        "缺少客户门店编码、本部或营业所时，异常说明清晰且字段显示“-”。",
        "责任人与列表状态一致；待复核时责任回到 POS 担当。",
        "驳回原因和营业处理说明在日志中均可追溯。",
    ])

    module_section(doc, "7", "质量检查—标准POS表", [
        ["入口", "门店视图或产品视图操作列“单据详情”"],
        ["粒度", "门店视图：门店；产品视图：当前产品明细"],
        ["状态来源", "内部状态映射为待通过 / 处理中 / 待复核 / 已通过"],
        ["原始文件", "使用标准POS来源文件名；缺失时按门店信息生成兜底文件名"],
    ], [
        "门店编码", "当前责任人", "最近操作人", "驳回原因", "营业处理说明", "置信度", "AI判断", "ACC", "本部", "营业所", "经销商"
    ], "仅产品视图追加：客户产品号、客户产品名称、客户条形码、好丽友产品编码、好丽友条形码、好丽友产品名称",
    [
        "生成标准POS：待通过，默认责任人为 POS 担当。",
        "驳回营业：处理中，记录 rejectNote / rejectedAt / rejectedBy。",
        "营业提交复核：待复核，记录 salesSubmitNote / salesSubmittedAt / salesSubmittedBy。",
        "POS通过：已通过，记录 approvedAt / approvedBy；日志说明当前单据已通过。",
    ], [
        "门店视图打开详情时不展示任意产品字段。",
        "产品视图必须从当前行 productIndex 取值，禁止默认取第一条产品。",
        "原始文件点击进入与当前门店/产品相关的数据预览。",
    ], [
        "不同视图字段粒度正确，产品详情不会串行。",
        "责任人、驳回原因和处理说明与工作流状态一致。",
        "已通过状态只读且日志完整。",
    ])

    module_section(doc, "8", "质量检查—异常数据", [
        ["入口", "门店视图或产品视图操作列“单据详情”"],
        ["粒度", "门店视图：异常门店批次代表记录；产品视图：单条异常产品"],
        ["状态来源", "待处理 / 处理中 / 待复核；单项解决状态可为已解决；整单通过后离开异常列表"],
        ["来源文件", "按同门店标准POS来源定位，无法定位时使用当前异常行兜底"],
    ], [
        "门店编码", "当前责任人", "最近操作人", "驳回原因", "营业处理说明", "异常类型", "AI判断", "ACC", "本部", "营业所", "经销商"
    ], "仅产品视图追加：客户产品号、客户产品名称、客户条形码、好丽友产品编码、好丽友条形码、好丽友产品名称、销售数量、销售金额、销售成本、零售单价",
    [
        "异常生成：异常处理/待处理，当前责任人为 POS 担当。",
        "POS驳回：异常驳回/处理中，展示驳回原因、时间、操作人。",
        "营业提交：提交复核/待复核，展示处理说明、时间、提交人。",
        "复核通过：异常复核/已通过，说明单据流入标准POS表。单项已解决但整单未通过时不得生成整单已通过日志。",
    ], [
        "门店视图详情应基于代表记录汇总批次状态；产品视图展示当前异常项。",
        "金额和数量保留 0；异常说明与 AI 判断不得互相覆盖。",
        "详情中不提供“异常项通过/整单通过”动作，动作留在列表。",
    ], [
        "门店与产品粒度切换后字段不串行。",
        "已解决、待整单通过属于内部/行级业务状态，不作为当前状态筛选枚举展示。",
        "整单通过后异常记录从异常列表退出，并可在标准POS/台账链路追溯。",
    ])

    module_section(doc, "9", "台账与汇总—标准POS明细", [
        ["入口", "标准POS明细操作列“单据详情”"],
        ["粒度", "产品销售明细"],
        ["记录定位", "month + storeCode + productCode/barcode + productName 组成行键"],
        ["状态", "固定“已入账”"],
        ["日志", "台账入库/已入账，说明标准POS通过后同步进入台账"],
    ], [
        "时间", "客户系统", "经销商", "客户门店号", "原始交易出码", "客户门店名称", "TEAM", "区域", "营业所", "ACC", "好丽友交易处编码", "好丽友交易处名称", "客户产品号", "客户产品名称", "客户条形码", "好丽友产品编码", "好丽友条形码", "好丽友产品名称", "销售数量", "销售金额", "成本", "零售单价"
    ], "可追加：入账时间、来源质检单号、最后操作人（当后端提供真实字段时）",
    [
        "仅展示台账入库节点；不得虚构前序节点时间。若后端返回完整生命周期，可按统一日志顺序补充。",
        "操作人优先 row.lastOperatorName，缺失显示“系统”。",
    ], [
        "只读查看；不在详情抽屉中编辑台账数值。",
        "原始文件可关联时支持预览；不可关联时使用兜底文件/数据预览。",
        "金额统一格式化，不拼接出“￥-”。",
    ], [
        "行键唯一定位当前产品记录，筛选或分组后仍打开正确行。",
        "客户与好丽友产品字段成对展示，编码不丢前导零。",
        "入账状态与日志一致。",
    ])

    add_heading(doc, "10. 状态与色彩映射", 1)
    add_table(doc, ["状态类别", "示例", "色彩/语义"], [
        ["异常/失败", "驳回、失败、异常", "红色；需要纠正或流程中断"],
        ["待办/风险", "待处理、待通过、暂存、未匹配、重复", "黄色；需要人工处理或确认"],
        ["进行中", "校验中、质检中、处理中、待复核", "蓝色；流程仍在推进"],
        ["成功/完成", "正常、已匹配、已完成、已通过、已入账", "绿色；节点已完成"],
        ["历史/无效", "历史版本、已归档", "灰色或中性；保留追溯，不参与当前处理"],
    ], [1800, 3000, 4560], font_size=9.1)

    add_heading(doc, "11. 前后端数据契约建议", 1)
    add_table(doc, ["字段组", "建议字段"], [
        ["详情上下文", "moduleCode, moduleName, currentNode, detailLevel, recordId, statusCode, statusText"],
        ["来源", "sourceEmailId, sourceAttachmentId/sourceAttIdx, sourceFileName, fileVersion, originalFileName"],
        ["业务字段", "moduleFields: [{key, label, value, valueType, group, order, visible}]"],
        ["拆分结果", "splitGroups: [{file, versions:[{version,state,uploadedAt,status,rows}]}]"],
        ["日志", "logs: [{nodeCode,nodeName,statusCode,statusText,action,time,operatorId,operatorName,tone}]"],
        ["能力", "permissions: {canViewDetail, canPreviewSource, canPreviewSplit}; 不返回无权访问的文件地址"],
    ], [1800, 7560], font_size=9.0)
    add_callout(doc, "建议", "生产实现应由后端返回标准化生命周期日志，前端只负责排序、转义和展示。当前原型按状态推导日志可保留作演示兜底，但不应成为审计数据来源。", "blue")

    add_heading(doc, "12. 联调与验收清单", 1)
    checks = [
        "权限：无查看或无单据详情权限时，按钮不可见且直接访问被拒绝。",
        "数据范围：跨营业所/团队记录不能通过修改 ID 打开。",
        "定位：父单据、附件、门店、产品四种粒度均打开正确记录。",
        "来源：文件名、版本、附件范围正确；驳回文件禁止预览。",
        "拆分：多门店、历史版本、驳回重传、重复处理均可追溯。",
        "字段：各模块字段顺序、空值、0 值、金额、编码格式符合本规格。",
        "日志：节点顺序、状态、原因、时间、操作人真实且一致。",
        "视图：门店视图不混入产品字段；产品视图不串行。",
        "安全：文件名、备注、日志动作等特殊字符均被转义，无 XSS。",
        "体验：抽屉滚动正常，关闭后保持列表筛选、分页和滚动位置。",
        "国际化：中文/韩文标题与固定文案由 i18n 管理；业务数据不做错误翻译。",
        "审计：历史日志不可被覆盖；缺失时间显示“-”而非伪造时间。",
    ]
    for c in checks:
        add_bullet(doc, "□ " + c)

    add_heading(doc, "附录 A：当前实现需补齐/统一项", 1)
    add_table(doc, ["优先级", "事项", "实现要求"], [
        ["P0", "模块字段未渲染", "在通用详情容器中新增“当前模块信息”区，消费 moduleFields，并按分组/顺序展示"],
        ["P0", "服务端鉴权", "详情与文件预览接口必须执行双权限和数据范围校验，不能只依赖前端按钮"],
        ["P1", "日志数据来源", "由后端提供真实节点日志；前端状态推导仅作为原型/兼容兜底"],
        ["P1", "金额/空值格式", "统一格式化器，正确处理 0、空值和币种，避免“￥-”"],
        ["P1", "归档权限", "为归档（非POS表）定义清晰的查看与详情权限，避免隐式放行"],
        ["P2", "国际化", "抽屉内固定中文文本全部进入 i18n；字段标签按模块语言包输出"],
        ["P2", "可访问性", "按钮提供 aria-label；抽屉支持焦点圈定、Esc 关闭与关闭后焦点归还"],
    ], [1000, 2650, 5710], font_size=8.9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
